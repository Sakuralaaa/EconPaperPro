# -*- coding: utf-8 -*-
"""
Word文档解析模块
使用 python-docx 解析 Word 文件，提取文本内容
"""

from typing import Dict, List

# 安全导入 python-docx
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    Document = None
    DOCX_AVAILABLE = False


class DocxParser:
    """
    Word 文档解析器
    
    使用示例:
        parser = DocxParser()
        text = parser.parse("paper.docx")
    """
    
    def parse(self, file_path: str) -> str:
        """
        解析 Word 文件，提取全部文本
        
        Args:
            file_path: Word 文件路径
            
        Returns:
            str: 提取的文本内容
        """
        if not DOCX_AVAILABLE or Document is None:
            raise RuntimeError("Word 解析需要安装 python-docx: pip install python-docx")
        
        try:
            doc = Document(file_path)
            paragraphs = [para.text for para in doc.paragraphs]
            return "\n".join(paragraphs)
            
        except Exception as e:
            raise RuntimeError(f"Word 文档解析失败: {str(e)}")
    
    def parse_with_structure(self, file_path: str) -> Dict:
        """
        解析 Word 文件并保留结构信息
        
        Args:
            file_path: Word 文件路径
            
        Returns:
            Dict: 包含段落和结构信息的字典
        """
        if not DOCX_AVAILABLE or Document is None:
            raise RuntimeError("Word 解析需要安装 python-docx: pip install python-docx")
        
        try:
            doc = Document(file_path)
            
            paragraphs = []
            for para in doc.paragraphs:
                style_name = para.style.name if para.style else "Normal"
                paragraphs.append({
                    "text": para.text,
                    "style": style_name,
                    "is_heading": style_name.startswith("Heading")
                })
            
            # 提取表格
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
            
            return {
                "paragraphs": paragraphs,
                "tables": tables,
                "full_text": "\n".join([p["text"] for p in paragraphs])
            }
            
        except Exception as e:
            raise RuntimeError(f"Word 结构化解析失败: {str(e)}")
    
    def parse_headings(self, file_path: str) -> List[Dict]:
        """
        提取 Word 文档的标题结构
        
        Args:
            file_path: Word 文件路径
            
        Returns:
            List[Dict]: 标题列表
        """
        if not DOCX_AVAILABLE or Document is None:
            raise RuntimeError("Word 解析需要安装 python-docx: pip install python-docx")
        
        try:
            doc = Document(file_path)
            
            headings = []
            for i, para in enumerate(doc.paragraphs):
                style_name = para.style.name if para.style else "Normal"
                if style_name.startswith("Heading"):
                    level = int(style_name.replace("Heading ", "")) if style_name != "Heading" else 1
                    headings.append({
                        "level": level,
                        "text": para.text,
                        "index": i
                    })
            
            return headings
            
        except Exception as e:
            raise RuntimeError(f"标题提取失败: {str(e)}")
    
    def parse_bytes(self, file_bytes: bytes) -> str:
        """
        从字节流解析 Word 文档
        
        Args:
            file_bytes: Word 文件字节流
            
        Returns:
            str: 提取的文本内容
        """
        if not DOCX_AVAILABLE or Document is None:
            raise RuntimeError("Word 解析需要安装 python-docx: pip install python-docx")
        
        try:
            from io import BytesIO
            doc = Document(BytesIO(file_bytes))
            paragraphs = [para.text for para in doc.paragraphs]
            return "\n".join(paragraphs)
            
        except Exception as e:
            raise RuntimeError(f"Word 字节流解析失败: {str(e)}")
