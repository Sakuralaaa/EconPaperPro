# -*- coding: utf-8 -*-
"""
EconPaper Pro - 原生 Tkinter GUI 应用 (v2.4流畅体验优化版)
- 修复UI卡顿问题
- 现代化界面设计
- 添加进度指示器
- 优化字体大小
- 分离API配置
- 模型拉取功能
- 首次使用引导
- 实时字数统计
- 关于页面
- v2.4新增:
  - 任务取消支持
  - 快捷键绑定
  - 工具提示
  - 状态栏
  - 通知横幅
  - 改进的动画效果
  - 更好的错误处理
"""

VERSION = "0.4.4"

import tkinter as tk
from tkinter import ttk, filedialog, scrolledtext
import threading
import sys
import os
from pathlib import Path
from typing import Optional, Callable, Dict, Any, List, Tuple
from datetime import datetime
import queue
import traceback

# 尝试导入 OpenAI
try:
    from openai import OpenAI
except ImportError:
    OpenAI = None

# 尝试导入 python-docx 用于 Word 导出
try:
    import docx  # type: ignore[import-untyped]
    HAS_DOCX = True
except ImportError:
    docx = None  # type: ignore[assignment]
    HAS_DOCX = False

# 导入自定义组件
from ui.components import (
    ModernStyle, Tooltip, StatusBar, AnimatedProgressBar,
    ModernButton, TaskManager, TextInputWithCount,
    TextOutputBox, NotificationBanner, KeyboardShortcuts,
    ConfirmDialog, DualOutputFrame, WorkflowConnector,
    PreciseProgressBar, StreamingTextOutput
)
from core.history import HistoryManager

# 确保模块路径正确
if getattr(sys, 'frozen', False):
    BASE_DIR = Path(sys.executable).parent
    INTERNAL_DIR = Path(getattr(sys, '_MEIPASS', BASE_DIR))
else:
    BASE_DIR = Path(__file__).parent.parent
    INTERNAL_DIR = BASE_DIR

sys.path.insert(0, str(BASE_DIR))
sys.path.insert(0, str(INTERNAL_DIR))



class EconPaperApp:
    """EconPaper Pro 主应用 - v2.4流畅体验优化版"""
    
    def __init__(self, root: tk.Tk):
        self.root = root
        self.root.title("📚 EconPaper Pro - 经管论文智能优化")
        self.root.geometry("1400x900")
        self.root.minsize(1100, 700)
        
        # 设置图标
        try:
            icon_path = BASE_DIR / "assets" / "icon.ico"
            if icon_path.exists():
                self.root.iconbitmap(str(icon_path))
        except Exception:
            pass
        
        # 配置样式
        self.style = ModernStyle.configure_styles(root)
        self.root.configure(bg=ModernStyle.BG_MAIN)
        
        # 任务队列
        self.update_queue = queue.Queue()
        
        # 任务管理器
        self.task_manager = TaskManager(self._safe_update)
        
        # 状态变量
        self.current_tab = tk.StringVar(value="diagnose")
        self.is_processing = False
        self.last_search_results = []  # 存储最近的搜索结果
        self.api_configured = False  # API是否已配置
        self.active_tasks = {}  # 活动任务跟踪
        
        # 历史记录管理器 - 必须在 _create_layout() 之前初始化
        # 因为 _create_template_selector() 需要使用 self.history
        self.history = HistoryManager()
        
        # 创建主布局
        self._create_layout()
        
        # 创建通知横幅
        self.notification = NotificationBanner(self.root)
        setattr(self.root, 'notification', self.notification) # 挂载到 root 方便组件访问，使用 setattr 绕过类型检查
        
        # 创建状态栏
        self._create_status_bar()
        
        # 绑定快捷键
        self.shortcuts = KeyboardShortcuts(self.root)
        self._bind_shortcuts()
        
        # 显示快捷键提示
        self.shortcuts.show_shortcut_hints(self.root)
        
        # 工作流连接器（管理页面间的数据流转）
        self.workflow = WorkflowConnector(self)
        
        # 启动UI更新循环
        self._process_queue()
        
        # 加载持久化偏好 (P2)
        self.root.after(100, self._load_ui_preferences)

        # 首次使用检查
        self.root.after(500, self._check_first_run)
        
        # 窗口关闭处理
        self.root.protocol("WM_DELETE_WINDOW", self._on_closing)
    
    def _create_status_bar(self):
        """创建状态栏"""
        self.status_bar = StatusBar(self.root)
        self.status_bar.set_info(f"v{VERSION}")
        self.status_bar.set_status("就绪", "normal")
    
    def _bind_shortcuts(self):
        """绑定快捷键"""
        # Ctrl+1-5 切换页面
        self.shortcuts.bind("<Control-Key-1>", lambda: self._show_page("diagnose"), "切换到论文诊断")
        self.shortcuts.bind("<Control-Key-2>", lambda: self._show_page("optimize"), "切换到深度优化")
        self.shortcuts.bind("<Control-Key-3>", lambda: self._show_page("dedup"), "切换到降重降AI")
        self.shortcuts.bind("<Control-Key-4>", lambda: self._show_page("search"), "切换到学术搜索")
        self.shortcuts.bind("<Control-Key-5>", lambda: self._show_page("revision"), "切换到退修助手")
        
        # Ctrl+, 打开设置
        self.shortcuts.bind("<Control-comma>", lambda: self._show_page("settings"), "打开系统设置")
        
        # Ctrl+S 保存（在设置页面）
        self.shortcuts.bind("<Control-s>", lambda: self._save_settings() if self.current_tab.get() == "settings" else None, "保存配置")
        
        # F1 帮助/关于
        self.shortcuts.bind("<F1>", self._show_about_dialog, "查看关于")
        
        # Escape 取消当前操作
        self.shortcuts.bind("<Escape>", self._on_escape, "取消当前任务")
    
    def _on_escape(self, event=None):
        """Escape 键处理 - 取消当前操作"""
        if self.is_processing:
            self.task_manager.cancel_all()
            self.notification.show("已请求取消所有正在运行的任务", "warning")
    
    def _on_closing(self):
        """窗口关闭处理"""
        # 检查是否有正在进行的任务
        if self.is_processing:
            if not ConfirmDialog.show(self.root, "确认退出", "有任务正在进行中，确定要退出吗？"):
                return
        self.root.destroy()
        
    def _process_queue(self):
        """处理队列中的UI更新任务"""
        try:
            while True:
                task = self.update_queue.get_nowait()
                if callable(task):
                    task()
        except queue.Empty:
            pass
        finally:
            self.root.after(50, self._process_queue)
    
    def _safe_update(self, func):
        """线程安全的UI更新"""
        self.update_queue.put(func)
        
    def _create_layout(self):
        """创建主布局"""
        main_container = tk.Frame(self.root, bg=ModernStyle.BG_MAIN)
        main_container.pack(fill=tk.BOTH, expand=True)
        
        self._create_sidebar(main_container)
        
        self.content_frame = tk.Frame(main_container, bg=ModernStyle.BG_MAIN)
        self.content_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        self.pages = {}
        self.progress_indicators = {}
        self.precise_progress = {} # P0 新增：精确进度条
        
        # 创建顶部工具栏 (P3)
        self._create_top_bar()
        
        self._create_diagnose_page()
        self._create_optimize_page()
        self._create_dedup_page()
        self._create_search_page()
        self._create_revision_page()
        self._create_history_page()
        self._create_settings_page()
        
        self._show_page("diagnose")
        
    def _create_sidebar(self, parent):
        """创建侧边栏 - 优化字体大小"""
        sidebar = tk.Frame(parent, bg=ModernStyle.BG_SIDEBAR, width=260)
        sidebar.pack(side=tk.LEFT, fill=tk.Y)
        sidebar.pack_propagate(False)
        
        separator = tk.Frame(sidebar, bg=ModernStyle.BORDER, width=1)
        separator.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Logo 区域
        logo_frame = tk.Frame(sidebar, bg=ModernStyle.BG_SIDEBAR)
        logo_frame.pack(fill=tk.X, pady=(35, 25), padx=25)
        
        title_container = tk.Frame(logo_frame, bg=ModernStyle.BG_SIDEBAR)
        title_container.pack(anchor="w")
        
        tk.Label(
            title_container,
            text="📚",
            font=(ModernStyle.FONT_FAMILY, 28),
            bg=ModernStyle.BG_SIDEBAR
        ).pack(side=tk.LEFT, padx=(0, 12))
        
        title_text = tk.Frame(title_container, bg=ModernStyle.BG_SIDEBAR)
        title_text.pack(side=tk.LEFT)
        
        tk.Label(
            title_text,
            text="EconPaper",
            font=(ModernStyle.FONT_FAMILY, 18, "bold"),
            bg=ModernStyle.BG_SIDEBAR,
            fg=ModernStyle.TEXT_PRIMARY
        ).pack(anchor="w")
        
        tk.Label(
            title_text,
            text="Pro",
            font=(ModernStyle.FONT_FAMILY, 18),
            bg=ModernStyle.BG_SIDEBAR,
            fg=ModernStyle.PRIMARY
        ).pack(anchor="w")
        
        tk.Label(
            logo_frame,
            text="经管学术论文智能助手",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_XS),
            bg=ModernStyle.BG_SIDEBAR,
            fg=ModernStyle.TEXT_MUTED
        ).pack(anchor="w", pady=(8, 0))
        
        sep = tk.Frame(sidebar, bg=ModernStyle.BORDER, height=1)
        sep.pack(fill=tk.X, padx=20, pady=10)
        
        # 导航菜单
        nav_frame = tk.Frame(sidebar, bg=ModernStyle.BG_SIDEBAR)
        nav_frame.pack(fill=tk.X, padx=12)
        
        self.nav_buttons = {}
        nav_items = [
            ("diagnose", "🔍", "论文诊断", "多维度分析评估"),
            ("optimize", "⚙️", "深度优化", "智能优化改写"),
            ("dedup", "🔧", "降重降AI", "降低重复率"),
            ("search", "🔎", "学术搜索", "文献检索"),
            ("revision", "📝", "退修助手", "回应审稿意见"),
            ("history", "🕒", "历史记录", "查看及恢复历史结果"),
        ]
        
        for page_id, icon, title, desc in nav_items:
            btn_frame = tk.Frame(nav_frame, bg=ModernStyle.BG_SIDEBAR, cursor="hand2")
            btn_frame.pack(fill=tk.X, pady=3)
            
            btn_inner = tk.Frame(btn_frame, bg=ModernStyle.BG_SIDEBAR, padx=15, pady=12)
            btn_inner.pack(fill=tk.X)
            
            # 添加左侧指示条
            indicator = tk.Frame(btn_frame, bg=ModernStyle.BG_SIDEBAR, width=4)
            indicator.place(relx=0, rely=0, relheight=1)
            
            tk.Label(
                btn_inner,
                text=icon,
                font=(ModernStyle.FONT_FAMILY, 16),
                bg=ModernStyle.BG_SIDEBAR
            ).pack(side=tk.LEFT)
            
            text_frame = tk.Frame(btn_inner, bg=ModernStyle.BG_SIDEBAR)
            text_frame.pack(side=tk.LEFT, padx=12)
            
            title_label = tk.Label(
                text_frame,
                text=title,
                font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_MD),
                bg=ModernStyle.BG_SIDEBAR,
                fg=ModernStyle.TEXT_PRIMARY
            )
            title_label.pack(anchor="w")
            
            desc_label = tk.Label(
                text_frame,
                text=desc,
                font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_XS),
                bg=ModernStyle.BG_SIDEBAR,
                fg=ModernStyle.TEXT_MUTED
            )
            desc_label.pack(anchor="w")
            
            self.nav_buttons[page_id] = {
                "frame": btn_frame,
                "inner": btn_inner,
                "indicator": indicator,
                "title": title_label,
                "desc": desc_label
            }
            
            for widget in [btn_frame, btn_inner, title_label, desc_label]:
                widget.bind("<Button-1>", lambda e, p=page_id: self._show_page(p))
                widget.bind("<Enter>", lambda e, p=page_id: self._on_nav_hover(p, True))
                widget.bind("<Leave>", lambda e, p=page_id: self._on_nav_hover(p, False))
        
        # 底部按钮区
        bottom_frame = tk.Frame(sidebar, bg=ModernStyle.BG_SIDEBAR)
        bottom_frame.pack(side=tk.BOTTOM, fill=tk.X, padx=12, pady=20)
        
        sep2 = tk.Frame(bottom_frame, bg=ModernStyle.BORDER, height=1)
        sep2.pack(fill=tk.X, pady=(0, 15))
        
        # 设置按钮
        settings_btn = tk.Frame(bottom_frame, bg=ModernStyle.BG_SIDEBAR, cursor="hand2")
        settings_btn.pack(fill=tk.X, pady=3)
        
        settings_inner = tk.Frame(settings_btn, bg=ModernStyle.BG_SIDEBAR, padx=15, pady=12)
        settings_inner.pack(fill=tk.X)
        
        settings_icon = tk.Label(
            settings_inner,
            text="⚙️",
            font=(ModernStyle.FONT_FAMILY, 16),
            bg=ModernStyle.BG_SIDEBAR,
            cursor="hand2"
        )
        settings_icon.pack(side=tk.LEFT)
        
        settings_text = tk.Label(
            settings_inner,
            text="系统设置",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_MD),
            bg=ModernStyle.BG_SIDEBAR,
            fg=ModernStyle.TEXT_PRIMARY,
            cursor="hand2"
        )
        settings_text.pack(side=tk.LEFT, padx=12)
        
        self.nav_buttons["settings"] = {
            "frame": settings_btn,
            "inner": settings_inner,
            "title": settings_text,
            "desc": None
        }
        
        def on_settings_click(e):
            self._show_page("settings")
        
        settings_btn.bind("<Button-1>", on_settings_click)
        settings_inner.bind("<Button-1>", on_settings_click)
        settings_icon.bind("<Button-1>", on_settings_click)
        settings_text.bind("<Button-1>", on_settings_click)
        
        # 关于按钮
        about_btn = tk.Frame(bottom_frame, bg=ModernStyle.BG_SIDEBAR, cursor="hand2")
        about_btn.pack(fill=tk.X, pady=3)
        
        about_inner = tk.Frame(about_btn, bg=ModernStyle.BG_SIDEBAR, padx=15, pady=10)
        about_inner.pack(fill=tk.X)
        
        about_icon = tk.Label(
            about_inner,
            text="ℹ️",
            font=(ModernStyle.FONT_FAMILY, 14),
            bg=ModernStyle.BG_SIDEBAR,
            cursor="hand2"
        )
        about_icon.pack(side=tk.LEFT)
        
        about_text = tk.Label(
            about_inner,
            text=f"关于 v{VERSION}",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_SM),
            bg=ModernStyle.BG_SIDEBAR,
            fg=ModernStyle.TEXT_MUTED,
            cursor="hand2"
        )
        about_text.pack(side=tk.LEFT, padx=12)
        
        def on_about_click(e):
            self._show_about_dialog()
        
        about_btn.bind("<Button-1>", on_about_click)
        about_inner.bind("<Button-1>", on_about_click)
        about_icon.bind("<Button-1>", on_about_click)
        about_text.bind("<Button-1>", on_about_click)

    def _on_nav_hover(self, page_id, is_enter):
        """导航悬停效果"""
        if page_id not in self.nav_buttons:
            return
            
        btn = self.nav_buttons[page_id]
        if self.current_tab.get() == page_id:
            return
            
        bg_color = ModernStyle.BG_HOVER if is_enter else ModernStyle.BG_SIDEBAR
        btn["frame"].config(bg=bg_color)
        btn["inner"].config(bg=bg_color)
        btn["title"].config(bg=bg_color)
        if btn["desc"]:
            btn["desc"].config(bg=bg_color)

    def _update_nav_style(self):
        """更新导航栏选中样式"""
        current = self.current_tab.get()
        for page_id, btn in self.nav_buttons.items():
            if page_id == current:
                bg_color = ModernStyle.PRIMARY_LIGHT
                btn["frame"].config(bg=bg_color)
                btn["inner"].config(bg=bg_color)
                btn["indicator"].config(bg=ModernStyle.PRIMARY)
                btn["title"].config(bg=bg_color, fg=ModernStyle.PRIMARY, font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_MD, "bold"))
                if btn["desc"]:
                    btn["desc"].config(bg=bg_color, fg=ModernStyle.PRIMARY)
            else:
                bg_color = ModernStyle.BG_SIDEBAR
                btn["frame"].config(bg=bg_color)
                btn["inner"].config(bg=bg_color)
                btn["indicator"].config(bg=bg_color)
                btn["title"].config(bg=bg_color, fg=ModernStyle.TEXT_PRIMARY, font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_MD))
                if btn["desc"]:
                    btn["desc"].config(bg=bg_color, fg=ModernStyle.TEXT_MUTED)
    
    def _show_page(self, page_id: str):
        """显示指定页面"""
        self.current_tab.set(page_id)
        self._update_nav_style()
        
        # 持久化当前页面偏好 (P2)
        self._save_ui_preference("last_page", page_id)
        
        for page in self.pages.values():
            page.pack_forget()
        
        if page_id in self.pages:
            self.pages[page_id].pack(fill=tk.BOTH, expand=True)
            
        # 更新状态栏
        page_names = {
            "diagnose": "论文诊断",
            "optimize": "深度优化",
            "dedup": "降重降AI",
            "search": "学术搜索",
            "revision": "退修助手",
            "history": "历史记录",
            "settings": "系统设置"
        }
        self.status_bar.set_status(f"当前页面: {page_names.get(page_id, page_id)}", "info")
    
    def _create_top_bar(self):
        """创建全局顶部工具栏 - 支持快速切换模型 (P3)"""
        self.top_bar = tk.Frame(self.content_frame, bg=ModernStyle.BG_MAIN, height=50)
        self.top_bar.pack(side=tk.TOP, fill=tk.X, padx=ModernStyle.PADDING_XL, pady=(ModernStyle.PADDING_MD, 0))
        self.top_bar.pack_propagate(False)

        # 页面标题占位 (由各页面动态更新或保持空白)
        self.top_title_var = tk.StringVar(value="")
        tk.Label(
            self.top_bar,
            textvariable=self.top_title_var,
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_LG, "bold"),
            bg=ModernStyle.BG_MAIN,
            fg=ModernStyle.TEXT_PRIMARY
        ).pack(side=tk.LEFT)

        # 右侧：模型快速选择
        model_frame = tk.Frame(self.top_bar, bg=ModernStyle.BG_MAIN)
        model_frame.pack(side=tk.RIGHT)

        tk.Label(
            model_frame,
            text="🤖 当前模型:",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_XS),
            bg=ModernStyle.BG_MAIN,
            fg=ModernStyle.TEXT_MUTED
        ).pack(side=tk.LEFT, padx=10)

        self.quick_model_var = tk.StringVar()
        self.quick_model_combo = ttk.Combobox(
            model_frame,
            textvariable=self.quick_model_var,
            values=["gpt-4o-mini", "gpt-4o", "deepseek-chat", "Qwen/Qwen2.5-72B-Instruct"],
            state="readonly",
            width=25,
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_XS)
        )
        self.quick_model_combo.pack(side=tk.LEFT)
        self.quick_model_combo.bind("<<ComboboxSelected>>", self._on_quick_model_change)
        
        # 初始加载模型名称
        self.root.after(1000, self._sync_quick_model_selector)

    def _sync_quick_model_selector(self):
        """同步快速选择器的模型名称"""
        try:
            from config.settings import settings
            if settings.llm_model:
                self.quick_model_var.set(settings.llm_model)
        except Exception:
            pass

    def _on_quick_model_change(self, event=None):
        """快速切换模型回调"""
        new_model = self.quick_model_var.get()
        try:
            from config.settings import settings
            settings.llm_model = new_model
            # 同时更新设置页面的显示
            if hasattr(self, 'setting_llm_model'):
                self.setting_llm_model.set(new_model)
            
            self.notification.show(f"模型已快速切换至: {new_model}", "success")
            # 保存到 .env 以便持久化
            self._save_settings_silent()
        except Exception as e:
            self.notification.show(f"切换失败: {e}", "error")

    def _save_settings_silent(self):
        """静默保存设置到 .env"""
        try:
            from config.settings import settings
            env_path = BASE_DIR / ".env"
            # 读取现有内容
            lines = []
            if env_path.exists():
                with open(env_path, "r", encoding="utf-8") as f:
                    for line in f:
                        if "LLM_MODEL=" in line:
                            lines.append(f"LLM_MODEL={settings.llm_model}\n")
                        else:
                            lines.append(line)
            
            with open(env_path, "w", encoding="utf-8") as f:
                f.writelines(lines)
        except Exception:
            pass

    def _create_page_header(self, parent, title, subtitle):
        """创建页面标题区域"""
        header = tk.Frame(parent, bg=ModernStyle.BG_MAIN)
        header.pack(fill=tk.X, padx=ModernStyle.PADDING_XL, pady=(ModernStyle.PADDING_XL, ModernStyle.PADDING_LG))
        
        tk.Label(
            header,
            text=title,
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_XXL, "bold"),
            bg=ModernStyle.BG_MAIN,
            fg=ModernStyle.TEXT_PRIMARY
        ).pack(anchor="w")
        
        tk.Label(
            header,
            text=subtitle,
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_MD),
            bg=ModernStyle.BG_MAIN,
            fg=ModernStyle.TEXT_SECONDARY
        ).pack(anchor="w", pady=(5, 0))
        
        return header
    
    
    def _create_diagnose_page(self):
        """创建论文诊断页面"""
        page = tk.Frame(self.content_frame, bg=ModernStyle.BG_MAIN)
        self.pages["diagnose"] = page
        
        self._create_page_header(page, "论文诊断", "多维度 AI 分析论文质量，提供改进建议")
        
        self.progress_indicators["diagnose"] = AnimatedProgressBar(page, "正在分析论文...")
        self.precise_progress["diagnose"] = PreciseProgressBar(page, "诊断进度")
        
        content = tk.Frame(page, bg=ModernStyle.BG_MAIN)
        content.pack(fill=tk.BOTH, expand=True, padx=ModernStyle.PADDING_XL, pady=(0, ModernStyle.PADDING_XL))
        
        paned = tk.PanedWindow(content, orient=tk.HORIZONTAL, bg=ModernStyle.BG_MAIN, sashwidth=8, sashrelief=tk.FLAT)
        paned.pack(fill=tk.BOTH, expand=True)
        
        # 左侧输入
        left_panel = tk.Frame(paned, bg=ModernStyle.BG_MAIN)
        
        toolbar = tk.Frame(left_panel, bg=ModernStyle.BG_MAIN)
        toolbar.pack(fill=tk.X, pady=(0, 15))
        
        upload_btn = ModernButton(
            toolbar,
            text="📁 选择文件",
            command=lambda: self._select_file("diagnose"),
            width=120,
            height=40,
            bg_color=ModernStyle.BG_SECONDARY,
            hover_color=ModernStyle.BG_HOVER,
            text_color=ModernStyle.TEXT_PRIMARY,
            tooltip="上传 PDF 或 Word 格式的论文"
        )
        upload_btn.pack(side=tk.LEFT)
        
        self.diag_file_label = tk.Label(
            toolbar,
            text="支持 PDF/Word 文档",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_SM),
            bg=ModernStyle.BG_MAIN,
            fg=ModernStyle.TEXT_MUTED,
            padx=15
        )
        self.diag_file_label.pack(side=tk.LEFT)
        
        tk.Label(
            left_panel,
            text="论文内容",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_MD, "bold"),
            bg=ModernStyle.BG_MAIN,
            fg=ModernStyle.TEXT_PRIMARY
        ).pack(anchor="w", pady=(0, 10))
        
        self.diag_input_comp = TextInputWithCount(left_panel, placeholder="在此粘贴论文内容或上传文件...", max_chars=30000)
        self.diag_input_comp.pack(fill=tk.BOTH, expand=True)
        self.diag_text = self.diag_input_comp.text # 保持兼容性
        
        btn_frame = tk.Frame(left_panel, bg=ModernStyle.BG_MAIN)
        btn_frame.pack(fill=tk.X, pady=(18, 0))
        
        ModernButton(
            btn_frame,
            text="开始诊断",
            command=self._run_diagnose,
            width=150,
            height=45,
            tooltip="启动 AI 论文质量评估"
        ).pack(side=tk.LEFT)
        
        # 模板选择 (P3)
        self._create_template_selector(btn_frame, "diagnose", self.diag_input_comp)
        
        ModernButton(
            btn_frame,
            text="💾 存为模板",
            command=lambda: self._save_as_template("diagnose", self.diag_input_comp),
            width=120,
            height=45,
            bg_color=ModernStyle.BG_SECONDARY,
            hover_color=ModernStyle.BG_HOVER,
            text_color=ModernStyle.TEXT_PRIMARY,
            tooltip="将当前输入内容保存为常用模板"
        ).pack(side=tk.LEFT, padx=15)
        
        # 添加文献推荐按钮
        ModernButton(
            btn_frame,
            text="📚 相关文献",
            command=self._recommend_literature,
            width=130,
            height=45,
            bg_color=ModernStyle.INFO,
            hover_color=ModernStyle.INFO,
            tooltip="基于论文内容推荐参考文献"
        ).pack(side=tk.LEFT, padx=15)
        
        paned.add(left_panel, minsize=350)
        
        # 右侧结果 - 使用双重输出框架
        right_panel = tk.Frame(paned, bg=ModernStyle.BG_MAIN)
        
        result_header = tk.Frame(right_panel, bg=ModernStyle.BG_MAIN)
        result_header.pack(fill=tk.X, pady=(0, 10))
        
        tk.Label(
            result_header,
            text="诊断结果",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_MD, "bold"),
            bg=ModernStyle.BG_MAIN,
            fg=ModernStyle.TEXT_PRIMARY
        ).pack(side=tk.LEFT)
        
        # 双重输出框架：分离内容与报告
        self.diag_dual_output = DualOutputFrame(
            right_panel,
            height=12,
            show_actions=True,
            on_send_to=lambda target, content, ctx=False: self.workflow.send_to_page(target, content, "diagnose", ctx)
        )
        self.diag_dual_output.pack(fill=tk.BOTH, expand=True)
        
        # 添加流转按钮
        self.diag_dual_output.add_flow_button("发送至优化", "optimize", "⚙️")
        self.diag_dual_output.add_flow_button("作为背景参考", "optimize", "📎", as_context=True)
        
        # 保持兼容性
        self.diag_result = self.diag_dual_output.content_output.text
        
        paned.add(right_panel, minsize=350)
        
        self.diag_file_path = None
        self.diag_file_paths = []
        
    def _create_optimize_page(self):
        """创建深度优化页面"""
        page = tk.Frame(self.content_frame, bg=ModernStyle.BG_MAIN)
        self.pages["optimize"] = page
        
        self._create_page_header(page, "深度优化", "针对不同阶段和期刊，对论文进行精细化打磨")
        
        self.progress_indicators["optimize"] = AnimatedProgressBar(page, "正在优化论文...")
        self.precise_progress["optimize"] = PreciseProgressBar(page, "优化进度")
        
        content = tk.Frame(page, bg=ModernStyle.BG_MAIN)
        content.pack(fill=tk.BOTH, expand=True, padx=ModernStyle.PADDING_XL, pady=(0, ModernStyle.PADDING_XL))
        
        # 左侧配置面板
        config_panel = tk.Frame(content, bg=ModernStyle.BG_SECONDARY, width=280)
        config_panel.pack(side=tk.LEFT, fill=tk.Y, padx=(0, 20))
        config_panel.pack_propagate(False)
        
        config_inner = tk.Frame(config_panel, bg=ModernStyle.BG_SECONDARY, padx=22, pady=22)
        config_inner.pack(fill=tk.BOTH, expand=True)
        
        # 优化阶段
        tk.Label(
            config_inner,
            text="优化阶段",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_MD, "bold"),
            bg=ModernStyle.BG_SECONDARY,
            fg=ModernStyle.TEXT_PRIMARY
        ).pack(anchor="w", pady=(0, 12))
        
        self.opt_stage = tk.StringVar(value="submission")
        stages = [
            ("初稿重构", "draft"),
            ("投稿优化", "submission"),
            ("退修回应", "revision"),
            ("终稿定稿", "final")
        ]
        
        for text, value in stages:
            rb = tk.Radiobutton(
                config_inner,
                text=text,
                variable=self.opt_stage,
                value=value,
                bg=ModernStyle.BG_SECONDARY,
                fg=ModernStyle.TEXT_PRIMARY,
                activebackground=ModernStyle.BG_SECONDARY,
                selectcolor=ModernStyle.BG_SECONDARY,
                font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_SM)
            )
            rb.pack(anchor="w", pady=3)
        
        # 目标期刊
        tk.Label(
            config_inner,
            text="目标期刊",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_MD, "bold"),
            bg=ModernStyle.BG_SECONDARY,
            fg=ModernStyle.TEXT_PRIMARY
        ).pack(anchor="w", pady=(22, 12))
        
        self.opt_journal = tk.StringVar(value="")
        journals = ["", "经济研究", "管理世界", "金融研究", "中国工业经济", "会计研究", "其他"]
        journal_combo = ttk.Combobox(
            config_inner,
            textvariable=self.opt_journal,
            values=journals,
            state="readonly",
            width=24,
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_SM)
        )
        journal_combo.pack(fill=tk.X)
        
        # 优化章节
        tk.Label(
            config_inner,
            text="优化章节",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_MD, "bold"),
            bg=ModernStyle.BG_SECONDARY,
            fg=ModernStyle.TEXT_PRIMARY
        ).pack(anchor="w", pady=(22, 12))
        
        sections = [
            ("标题", "title"),
            ("摘要", "abstract"),
            ("引言", "introduction"),
            ("文献综述", "literature"),
            ("理论假设", "theory"),
            ("研究方法", "methodology"),
            ("实证结果", "results"),
            ("结论", "conclusion")
        ]
        
        self.opt_sections = {}
        for text, value in sections:
            var = tk.BooleanVar(value=value in ["abstract", "introduction"])
            self.opt_sections[value] = var
            cb = tk.Checkbutton(
                config_inner,
                text=text,
                variable=var,
                bg=ModernStyle.BG_SECONDARY,
                fg=ModernStyle.TEXT_PRIMARY,
                activebackground=ModernStyle.BG_SECONDARY,
                selectcolor=ModernStyle.BG_SECONDARY,
                font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_SM)
            )
            cb.pack(anchor="w", pady=2)
        
        # 文件上传
        tk.Label(
            config_inner,
            text="上传文件",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_MD, "bold"),
            bg=ModernStyle.BG_SECONDARY,
            fg=ModernStyle.TEXT_PRIMARY
        ).pack(anchor="w", pady=(22, 12))
        
        ModernButton(
            config_inner,
            text="📁 选择文件",
            command=lambda: self._select_file("optimize"),
            width=220,
            height=40,
            bg_color=ModernStyle.BG_MAIN,
            hover_color=ModernStyle.BG_HOVER,
            text_color=ModernStyle.TEXT_PRIMARY,
            tooltip="上传需要优化的论文文件"
        ).pack(fill=tk.X)
        
        self.opt_file_label = tk.Label(
            config_inner,
            text="未选择文件",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_XS),
            bg=ModernStyle.BG_SECONDARY,
            fg=ModernStyle.TEXT_MUTED,
            wraplength=220
        )
        self.opt_file_label.pack(pady=8)
        
        ModernButton(
            config_inner,
            text="开始优化",
            command=self._run_optimize,
            width=220,
            height=45,
            tooltip="根据配置启动深度优化"
        ).pack(side=tk.BOTTOM, pady=12)
        
        # 右侧编辑区 - 必须先创建以便模板选择器引用 opt_input_comp
        right_panel = tk.Frame(content, bg=ModernStyle.BG_MAIN)
        right_panel.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True)
        
        # 新增：可折叠的参考背景/审稿意见区域 (P1)
        self.opt_context_visible = tk.BooleanVar(value=False)
        context_header = tk.Frame(right_panel, bg=ModernStyle.BG_MAIN)
        context_header.pack(fill=tk.X, pady=(0, 5))
        
        tk.Label(
            context_header,
            text="📎 参考背景 / 审稿意见 (可选)",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_SM, "bold"),
            bg=ModernStyle.BG_MAIN,
            fg=ModernStyle.TEXT_SECONDARY
        ).pack(side=tk.LEFT)
        
        self.context_toggle_btn = tk.Label(
            context_header,
            text="[ 展开 + ]",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_XS),
            bg=ModernStyle.BG_MAIN,
            fg=ModernStyle.PRIMARY,
            cursor="hand2"
        )
        self.context_toggle_btn.pack(side=tk.LEFT, padx=10)
        self.context_toggle_btn.bind("<Button-1>", lambda e: self._toggle_opt_context())
        
        self.opt_context_frame = tk.Frame(right_panel, bg=ModernStyle.BG_MAIN)
        # 初始不展示
        
        self.opt_context_input = TextInputWithCount(self.opt_context_frame, height=5, placeholder="在此粘贴参考文献摘要或审稿意见，AI 将以此作为优化参考...", show_count=False)
        self.opt_context_input.pack(fill=tk.X, pady=(0, 10))
        
        tk.Label(
            right_panel,
            text="论文内容",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_MD, "bold"),
            bg=ModernStyle.BG_MAIN
        ).pack(anchor="w", pady=(5, 8))
        
        self.opt_input_comp = TextInputWithCount(right_panel, height=10, placeholder="在此粘贴论文内容...", max_chars=15000)
        self.opt_input_comp.pack(fill=tk.BOTH, expand=True, pady=(0, 18))
        self.opt_input = self.opt_input_comp.text
        
        # 模板选择 (P3) - 现在 opt_input_comp 已创建，可以安全引用
        self._create_template_selector(config_inner, "optimize", self.opt_input_comp)
        
        ModernButton(
            config_inner,
            text="💾 存为模板",
            command=lambda: self._save_as_template("optimize", self.opt_input_comp),
            width=220,
            height=40,
            bg_color=ModernStyle.BG_MAIN,
            hover_color=ModernStyle.BG_HOVER,
            text_color=ModernStyle.TEXT_PRIMARY,
            tooltip="保存当前优化指令为模板"
        ).pack(fill=tk.X, pady=8)
        
        tk.Label(
            right_panel,
            text="优化结果",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_MD, "bold"),
            bg=ModernStyle.BG_MAIN
        ).pack(anchor="w", pady=(0, 8))
        
        # 双重输出框架：分离优化后内容与修改报告
        self.opt_dual_output = DualOutputFrame(
            right_panel,
            height=10,
            show_actions=True,
            on_send_to=lambda target, content, ctx=False: self.workflow.send_to_page(target, content, "optimize", ctx)
        )
        self.opt_dual_output.pack(fill=tk.BOTH, expand=True)
        
        # 添加流转按钮
        self.opt_dual_output.add_flow_button("发送至降重", "dedup", "🔧")
        self.opt_dual_output.add_flow_button("作为背景参考", "revision", "📎", as_context=True)
        
        # 保持兼容性
        self.opt_output = self.opt_dual_output.content_output.text
        
        self.opt_file_path = None
        self.opt_file_paths = []

    def _toggle_opt_context(self, show: Optional[bool] = None):
        """切换优化页面的背景参考区域显示状态"""
        if show is not None:
            self.opt_context_visible.set(show)
        else:
            self.opt_context_visible.set(not self.opt_context_visible.get())
        
        if self.opt_context_visible.get():
            # 修复: 避免使用 pack(after=...) 参数，改用显式的 pack 顺序控制
            # 先隐藏后续内容，插入 context_frame，再显示后续内容
            self.opt_context_frame.pack(fill=tk.X, pady=(0, 10))
            # 确保 context_frame 在正确位置（toggle_btn 之后）
            try:
                self.opt_context_frame.lift()
            except Exception:
                pass
            self.context_toggle_btn.config(text="[ 折叠 - ]")
        else:
            self.opt_context_frame.pack_forget()
            self.context_toggle_btn.config(text="[ 展开 + ]")
        
    def _create_dedup_page(self):
        """创建降重降AI页面"""
        page = tk.Frame(self.content_frame, bg=ModernStyle.BG_MAIN)
        self.pages["dedup"] = page
        
        self._create_page_header(page, "降重与降AI", "智能改写文本，降低重复率与AI检测痕迹")
        
        self.progress_indicators["dedup"] = AnimatedProgressBar(page, "正在处理文本...")
        self.precise_progress["dedup"] = PreciseProgressBar(page, "处理进度")
        
        content = tk.Frame(page, bg=ModernStyle.BG_MAIN)
        content.pack(fill=tk.BOTH, expand=True, padx=ModernStyle.PADDING_XL, pady=(0, ModernStyle.PADDING_XL))
        
        # 参数栏
        params_frame = tk.Frame(content, bg=ModernStyle.BG_SECONDARY, padx=22, pady=18)
        params_frame.pack(fill=tk.X, pady=(0, 22))
        
        tk.Label(
            params_frame,
            text="处理强度:",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_MD),
            bg=ModernStyle.BG_SECONDARY,
            fg=ModernStyle.TEXT_PRIMARY
        ).pack(side=tk.LEFT)
        
        self.dedup_strength = tk.Scale(
            params_frame,
            from_=1, to=5,
            orient=tk.HORIZONTAL,
            length=160,
            bg=ModernStyle.BG_SECONDARY,
            fg=ModernStyle.TEXT_PRIMARY,
            highlightthickness=0,
            troughcolor=ModernStyle.BORDER,
            activebackground=ModernStyle.PRIMARY,
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_SM)
        )
        self.dedup_strength.set(3)
        self.dedup_strength.pack(side=tk.LEFT, padx=12)
        
        tk.Label(
            params_frame,
            text="1轻度 ←→ 5深度",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_XS),
            bg=ModernStyle.BG_SECONDARY,
            fg=ModernStyle.TEXT_MUTED
        ).pack(side=tk.LEFT, padx=8)
        
        tk.Label(
            params_frame,
            text="保留术语:",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_MD),
            bg=ModernStyle.BG_SECONDARY,
            fg=ModernStyle.TEXT_PRIMARY
        ).pack(side=tk.LEFT, padx=(35, 0))
        
        self.dedup_terms = tk.Entry(
            params_frame,
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_MD),
            width=32,
            bg=ModernStyle.BG_MAIN,
            relief="flat"
        )
        self.dedup_terms.pack(side=tk.LEFT, padx=12, ipady=6)
        self.dedup_terms.insert(0, "用逗号分隔，如: DID, PSM")
        self.dedup_terms.bind("<FocusIn>", lambda e: self.dedup_terms.delete(0, tk.END) if "逗号分隔" in self.dedup_terms.get() else None)
        
        # 文件上传 (P3)
        ModernButton(
            params_frame,
            text="📁 选择文件",
            command=lambda: self._select_file("dedup"),
            width=120,
            height=36,
            bg_color=ModernStyle.BG_MAIN,
            hover_color=ModernStyle.BG_HOVER,
            text_color=ModernStyle.TEXT_PRIMARY,
            tooltip="上传需要处理的文档 (支持多选)"
        ).pack(side=tk.RIGHT, padx=10)
        
        self.dedup_file_label = tk.Label(
            params_frame,
            text="",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_XS),
            bg=ModernStyle.BG_SECONDARY,
            fg=ModernStyle.SUCCESS
        )
        self.dedup_file_label.pack(side=tk.RIGHT, padx=5)
        
        # 文本区域
        text_frame = tk.Frame(content, bg=ModernStyle.BG_MAIN)
        text_frame.pack(fill=tk.BOTH, expand=True)
        
        paned = tk.PanedWindow(text_frame, orient=tk.HORIZONTAL, bg=ModernStyle.BG_MAIN, sashwidth=8)
        paned.pack(fill=tk.BOTH, expand=True)
        
        # 左侧输入
        left_panel = tk.Frame(paned, bg=ModernStyle.BG_MAIN)
        
        tk.Label(
            left_panel,
            text="原始文本",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_MD, "bold"),
            bg=ModernStyle.BG_MAIN
        ).pack(anchor="w", pady=(0, 8))
        
        self.dedup_input_comp = TextInputWithCount(left_panel, placeholder="在此粘贴需要降重或降AI的文本...", max_chars=8000)
        self.dedup_input_comp.pack(fill=tk.BOTH, expand=True)
        self.dedup_input = self.dedup_input_comp.text
        
        paned.add(left_panel, minsize=350)
        
        # 中间按钮
        mid_panel = tk.Frame(paned, bg=ModernStyle.BG_MAIN, width=160)
        mid_panel.pack_propagate(False)
        
        btn_container = tk.Frame(mid_panel, bg=ModernStyle.BG_MAIN)
        btn_container.place(relx=0.5, rely=0.5, anchor="center")
        
        buttons = [
            ("📉 智能降重", self._run_dedup, ModernStyle.PRIMARY),
            ("🤖 降AI痕迹", self._run_deai, ModernStyle.INFO),
            ("⚡ 深度全改", self._run_both_dedup, ModernStyle.SUCCESS)
        ]
        
        for text, cmd, color in buttons:
            ModernButton(
                btn_container,
                text=text,
                command=cmd,
                width=130,
                height=42,
                bg_color=color,
                hover_color=color
            ).pack(pady=10)
        
        paned.add(mid_panel, minsize=160)
        
        # 右侧输出 - 使用双重输出框架
        right_panel = tk.Frame(paned, bg=ModernStyle.BG_MAIN)
        
        dedup_result_header = tk.Frame(right_panel, bg=ModernStyle.BG_MAIN)
        dedup_result_header.pack(fill=tk.X, pady=(0, 8))
        
        tk.Label(
            dedup_result_header,
            text="改写结果",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_MD, "bold"),
            bg=ModernStyle.BG_MAIN
        ).pack(side=tk.LEFT)
        
        # 双重输出框架：分离改写结果与降重报告
        self.dedup_dual_output = DualOutputFrame(
            right_panel,
            height=12,
            show_actions=True,
            on_send_to=lambda target, content, ctx=False: self.workflow.send_to_page(target, content, "dedup", ctx)
        )
        self.dedup_dual_output.pack(fill=tk.BOTH, expand=True)
        
        self.dedup_dual_output.add_flow_button("发送至优化", "optimize", "⚙️")
        
        # 保持兼容性
        self.dedup_output = self.dedup_dual_output.content_output.text
        
        paned.add(right_panel, minsize=350)
        
        self.dedup_file_paths = []
        
    def _create_search_page(self):
        """创建学术搜索页面 - v2.0 多数据源学术检索"""
        page = tk.Frame(self.content_frame, bg=ModernStyle.BG_MAIN)
        self.pages["search"] = page
        
        self._create_page_header(page, "学术搜索", "中英文学术文献检索 - 支持多数据源")
        
        self.progress_indicators["search"] = AnimatedProgressBar(page, "正在搜索文献...")
        self.precise_progress["search"] = PreciseProgressBar(page, "搜索进度")
        
        content = tk.Frame(page, bg=ModernStyle.BG_MAIN)
        content.pack(fill=tk.BOTH, expand=True, padx=ModernStyle.PADDING_XL, pady=(0, ModernStyle.PADDING_XL))
        
        # 搜索栏
        search_frame = tk.Frame(content, bg=ModernStyle.BG_SECONDARY, padx=22, pady=18)
        search_frame.pack(fill=tk.X, pady=(0, 15))
        
        tk.Label(
            search_frame,
            text="🔍",
            font=(ModernStyle.FONT_FAMILY, 18),
            bg=ModernStyle.BG_SECONDARY
        ).pack(side=tk.LEFT, padx=(0, 12))
        
        self.search_query = tk.Entry(
            search_frame,
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_LG),
            bg=ModernStyle.BG_SECONDARY,
            fg=ModernStyle.TEXT_PRIMARY,
            relief="flat",
            width=40
        )
        self.search_query.pack(side=tk.LEFT, fill=tk.X, expand=True, ipady=6)
        self.search_query.insert(0, "digital economy innovation")
        
        # 绑定回车键搜索
        self.search_query.bind("<Return>", lambda e: self._run_search())
        
        # AI辅助按钮
        ModernButton(
            search_frame,
            text="🤖 AI扩展关键词",
            command=self._ai_expand_keywords,
            width=140,
            height=40,
            bg_color=ModernStyle.INFO,
            hover_color=ModernStyle.INFO
        ).pack(side=tk.LEFT, padx=12)
        
        # 数据源选择 - 中英文双语支持
        self.search_source = tk.StringVar(value="英文文献")
        source_combo = ttk.Combobox(
            search_frame,
            textvariable=self.search_source,
            values=["英文文献", "中文文献", "Semantic Scholar", "OpenAlex", "百度学术"],
            state="readonly",
            width=14,
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_SM)
        )
        source_combo.pack(side=tk.LEFT, padx=12)
        
        ModernButton(
            search_frame,
            text="搜索",
            command=self._run_search,
            width=100,
            height=40
        ).pack(side=tk.LEFT)
        
        # 筛选选项行1
        filter_frame = tk.Frame(content, bg=ModernStyle.BG_SECONDARY, padx=22, pady=12)
        filter_frame.pack(fill=tk.X, pady=(0, 8))
        
        tk.Label(
            filter_frame,
            text="结果数量:",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_SM),
            bg=ModernStyle.BG_SECONDARY,
            fg=ModernStyle.TEXT_PRIMARY
        ).pack(side=tk.LEFT, padx=(0, 8))
        
        self.search_limit = tk.Scale(
            filter_frame,
            from_=5, to=50,
            orient=tk.HORIZONTAL,
            length=100,
            bg=ModernStyle.BG_SECONDARY,
            fg=ModernStyle.TEXT_PRIMARY,
            highlightthickness=0,
            troughcolor=ModernStyle.BORDER,
            activebackground=ModernStyle.PRIMARY,
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_XS)
        )
        self.search_limit.set(15)
        self.search_limit.pack(side=tk.LEFT, padx=8)
        
        # 年份筛选
        tk.Label(
            filter_frame,
            text="起始年份:",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_SM),
            bg=ModernStyle.BG_SECONDARY,
            fg=ModernStyle.TEXT_PRIMARY
        ).pack(side=tk.LEFT, padx=(15, 8))
        
        self.search_year_from = tk.Entry(
            filter_frame,
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_SM),
            width=6,
            bg=ModernStyle.BG_MAIN,
            relief="flat"
        )
        self.search_year_from.pack(side=tk.LEFT, ipady=4)
        self.search_year_from.insert(0, "2020")
        
        self.enable_ai_filter = tk.BooleanVar(value=False)
        tk.Checkbutton(
            filter_frame,
            text="✨ AI智能筛选",
            variable=self.enable_ai_filter,
            bg=ModernStyle.BG_SECONDARY,
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_SM)
        ).pack(side=tk.LEFT, padx=(20, 8))
        
        tk.Label(
            filter_frame,
            text="💡 英文文献用英文关键词，中文文献用中文关键词",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_XS),
            bg=ModernStyle.BG_SECONDARY,
            fg=ModernStyle.TEXT_MUTED
        ).pack(side=tk.RIGHT, padx=12)
        
        # 期刊级别筛选行2
        quality_frame = tk.Frame(content, bg=ModernStyle.BG_SECONDARY, padx=22, pady=12)
        quality_frame.pack(fill=tk.X, pady=(0, 15))
        
        tk.Label(
            quality_frame,
            text="📊 期刊级别筛选:",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_SM, "bold"),
            bg=ModernStyle.BG_SECONDARY,
            fg=ModernStyle.TEXT_PRIMARY
        ).pack(side=tk.LEFT, padx=(0, 12))
        
        self.filter_cssci = tk.BooleanVar(value=False)
        tk.Checkbutton(
            quality_frame,
            text="仅CSSCI/北核",
            variable=self.filter_cssci,
            bg=ModernStyle.BG_SECONDARY,
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_SM)
        ).pack(side=tk.LEFT, padx=8)
        
        self.filter_ssci = tk.BooleanVar(value=False)
        tk.Checkbutton(
            quality_frame,
            text="仅SSCI Q1/Q2",
            variable=self.filter_ssci,
            bg=ModernStyle.BG_SECONDARY,
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_SM)
        ).pack(side=tk.LEFT, padx=8)
        
        self.show_rank_info = tk.BooleanVar(value=True)
        tk.Checkbutton(
            quality_frame,
            text="显示期刊级别",
            variable=self.show_rank_info,
            bg=ModernStyle.BG_SECONDARY,
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_SM)
        ).pack(side=tk.LEFT, padx=8)
        
        tk.Label(
            quality_frame,
            text="(基于内置期刊数据库，覆盖经管类核心期刊)",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_XS),
            bg=ModernStyle.BG_SECONDARY,
            fg=ModernStyle.TEXT_MUTED
        ).pack(side=tk.LEFT, padx=12)
        
        # 功能按钮区
        action_frame = tk.Frame(content, bg=ModernStyle.BG_SECONDARY, padx=22, pady=12)
        action_frame.pack(fill=tk.X, pady=(0, 15))
        
        ModernButton(
            action_frame,
            text="📝 生成文献综述",
            command=self._generate_literature_review,
            width=140,
            height=38,
            bg_color=ModernStyle.SUCCESS,
            hover_color=ModernStyle.SUCCESS
        ).pack(side=tk.LEFT, padx=8)
        
        ModernButton(
            action_frame,
            text="📋 生成引用格式",
            command=self._generate_citations,
            width=140,
            height=38,
            bg_color=ModernStyle.INFO,
            hover_color=ModernStyle.INFO
        ).pack(side=tk.LEFT, padx=8)
        
        ModernButton(
            action_frame,
            text="📥 导出结果",
            command=lambda: self._export_result(self.search_result.get("1.0", tk.END), "搜索结果"),
            width=110,
            height=38,
            bg_color=ModernStyle.TEXT_SECONDARY,
            hover_color=ModernStyle.TEXT_SECONDARY
        ).pack(side=tk.LEFT, padx=8)
        
        tk.Label(
            action_frame,
            text="📊 英文：Semantic Scholar + OpenAlex | 中文：百度学术 + 万方数据",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_XS),
            bg=ModernStyle.BG_SECONDARY,
            fg=ModernStyle.TEXT_MUTED
        ).pack(side=tk.RIGHT, padx=12)
        
        # 结果区 - 升级为 DualOutputFrame (P0)
        result_header = tk.Frame(content, bg=ModernStyle.BG_MAIN)
        result_header.pack(fill=tk.X, pady=(0, 10))
        
        tk.Label(
            result_header,
            text="搜索结果",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_MD, "bold"),
            bg=ModernStyle.BG_MAIN
        ).pack(side=tk.LEFT)
        
        self.search_status_label = tk.Label(
            result_header,
            text="",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_XS),
            bg=ModernStyle.BG_MAIN,
            fg=ModernStyle.TEXT_MUTED
        )
        self.search_status_label.pack(side=tk.RIGHT)
        
        # 使用双重输出框架：分离搜索结果与 AI 分析
        self.search_dual_output = DualOutputFrame(
            content,
            height=15,
            show_actions=True,
            on_send_to=lambda target, content, ctx=False: self.workflow.send_to_page(target, content, "search", ctx)
        )
        self.search_dual_output.pack(fill=tk.BOTH, expand=True)
        
        # 添加流转按钮
        self.search_dual_output.add_flow_button("作为参考背景", "optimize", "📎", as_context=True)
        self.search_dual_output.add_flow_button("发送至退修", "revision", "📝", as_context=True)
        
        # 保持兼容性
        self.search_result = self.search_dual_output.content_output.text
        
    def _create_revision_page(self):
        """创建退修助手页面"""
        page = tk.Frame(self.content_frame, bg=ModernStyle.BG_MAIN)
        self.pages["revision"] = page
        
        self._create_page_header(page, "退修助手", "智能解析审稿意见，生成逐条回应策略")
        
        self.progress_indicators["revision"] = AnimatedProgressBar(page, "正在分析审稿意见...")
        self.precise_progress["revision"] = PreciseProgressBar(page, "分析进度")
        
        content = tk.Frame(page, bg=ModernStyle.BG_MAIN)
        content.pack(fill=tk.BOTH, expand=True, padx=ModernStyle.PADDING_XL, pady=(0, ModernStyle.PADDING_XL))
        
        paned = tk.PanedWindow(content, orient=tk.HORIZONTAL, bg=ModernStyle.BG_MAIN, sashwidth=8)
        paned.pack(fill=tk.BOTH, expand=True)
        
        # 左侧输入
        left_panel = tk.Frame(paned, bg=ModernStyle.BG_MAIN)
        
        tk.Label(
            left_panel,
            text="审稿意见",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_MD, "bold"),
            bg=ModernStyle.BG_MAIN
        ).pack(anchor="w", pady=(0, 8))
        
        self.rev_comments_comp = TextInputWithCount(left_panel, height=12, placeholder="在此粘贴审稿意见...")
        self.rev_comments_comp.pack(fill=tk.BOTH, expand=True, pady=(0, 18))
        self.rev_comments = self.rev_comments_comp.text
        
        tk.Label(
            left_panel,
            text="论文摘要（可选）",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_MD, "bold"),
            bg=ModernStyle.BG_MAIN
        ).pack(anchor="w", pady=(0, 8))
        
        self.rev_summary_comp = TextInputWithCount(left_panel, height=6, placeholder="在此粘贴论文摘要，有助于AI理解上下文...")
        self.rev_summary_comp.pack(fill=tk.X, pady=(0, 18))
        self.rev_summary = self.rev_summary_comp.text
        
        rev_btn_frame = tk.Frame(left_panel, bg=ModernStyle.BG_MAIN)
        rev_btn_frame.pack(fill=tk.X)
        
        ModernButton(
            rev_btn_frame,
            text="生成回应策略",
            command=self._run_revision,
            width=180,
            height=45
        ).pack(side=tk.LEFT)
        
        # 模板选择 (P3)
        self._create_template_selector(rev_btn_frame, "revision", self.rev_comments_comp)
        
        ModernButton(
            rev_btn_frame,
            text="💾 存为模板",
            command=lambda: self._save_as_template("revision", self.rev_comments_comp),
            width=120,
            height=45,
            bg_color=ModernStyle.BG_SECONDARY,
            hover_color=ModernStyle.BG_HOVER,
            text_color=ModernStyle.TEXT_PRIMARY,
            tooltip="保存当前审稿意见或回应策略为模板"
        ).pack(side=tk.LEFT, padx=15)
        
        # 添加文献支撑按钮
        ModernButton(
            rev_btn_frame,
            text="📚 找支撑文献",
            command=self._find_supporting_literature,
            width=140,
            height=45,
            bg_color=ModernStyle.INFO,
            hover_color=ModernStyle.INFO
        ).pack(side=tk.LEFT, padx=15)
        
        paned.add(left_panel, minsize=400)
        
        # 右侧结果 - 使用双重输出框架
        right_panel = tk.Frame(paned, bg=ModernStyle.BG_MAIN)
        
        tk.Label(
            right_panel,
            text="回应建议",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_MD, "bold"),
            bg=ModernStyle.BG_MAIN
        ).pack(anchor="w", pady=(0, 10))
        
        # 双重输出框架：分离回应策略与详细分析
        self.rev_dual_output = DualOutputFrame(
            right_panel,
            height=12,
            show_actions=True,
            on_send_to=lambda target, content, ctx=False: self.workflow.send_to_page(target, content, "revision", ctx)
        )
        self.rev_dual_output.pack(fill=tk.BOTH, expand=True)
        
        # 添加流转按钮
        self.rev_dual_output.add_flow_button("查找文献", "search", "🔎")
        self.rev_dual_output.add_flow_button("发送至优化", "optimize", "⚙️")
        
        # 保持兼容性
        self.rev_output = self.rev_dual_output.content_output.text
        
        paned.add(right_panel, minsize=400)
        
    def _create_history_page(self):
        """创建历史记录页面"""
        page = tk.Frame(self.content_frame, bg=ModernStyle.BG_MAIN)
        self.pages["history"] = page
        
        self._create_page_header(page, "历史记录", "查看并恢复之前的 AI 生成结果及分析报告")
        
        content = tk.Frame(page, bg=ModernStyle.BG_MAIN)
        content.pack(fill=tk.BOTH, expand=True, padx=ModernStyle.PADDING_XL, pady=(0, ModernStyle.PADDING_XL))
        
        paned = tk.PanedWindow(content, orient=tk.HORIZONTAL, bg=ModernStyle.BG_MAIN, sashwidth=8, sashrelief=tk.FLAT)
        paned.pack(fill=tk.BOTH, expand=True)
        
        # 左侧列表
        left_panel = tk.Frame(paned, bg=ModernStyle.BG_MAIN)
        
        toolbar = tk.Frame(left_panel, bg=ModernStyle.BG_MAIN)
        toolbar.pack(fill=tk.X, pady=(0, 10))
        
        ModernButton(
            toolbar,
            text="🔄 刷新列表",
            command=self._refresh_history,
            width=100,
            height=32,
            bg_color=ModernStyle.BG_SECONDARY,
            hover_color=ModernStyle.BG_HOVER,
            text_color=ModernStyle.TEXT_PRIMARY
        ).pack(side=tk.LEFT)
        
        ModernButton(
            toolbar,
            text="🗑️ 清空历史",
            command=self._clear_all_history,
            width=100,
            height=32,
            bg_color=ModernStyle.BG_MAIN,
            hover_color=ModernStyle.ERROR_LIGHT,
            text_color=ModernStyle.TEXT_SECONDARY
        ).pack(side=tk.RIGHT)
        
        # Treeview
        columns = ("time", "type", "preview")
        self.history_tree = ttk.Treeview(
            left_panel,
            columns=columns,
            show="headings",
            selectmode="browse"
        )
        
        self.history_tree.heading("time", text="时间")
        self.history_tree.heading("type", text="类型")
        self.history_tree.heading("preview", text="内容预览")
        
        self.history_tree.column("time", width=150, minwidth=150)
        self.history_tree.column("type", width=100, minwidth=100)
        self.history_tree.column("preview", width=300)
        
        scrollbar = ttk.Scrollbar(left_panel, orient="vertical", command=self.history_tree.yview)
        self.history_tree.configure(yscrollcommand=scrollbar.set)
        
        self.history_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        self.history_tree.bind("<<TreeviewSelect>>", self._on_history_select)
        
        paned.add(left_panel, minsize=400)
        
        # 右侧详情
        right_panel = tk.Frame(paned, bg=ModernStyle.BG_MAIN)
        
        detail_header = tk.Frame(right_panel, bg=ModernStyle.BG_MAIN)
        detail_header.pack(fill=tk.X, pady=(0, 10))
        
        tk.Label(
            detail_header,
            text="记录详情",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_MD, "bold"),
            bg=ModernStyle.BG_MAIN
        ).pack(side=tk.LEFT)
        
        self.restore_btn = ModernButton(
            detail_header,
            text="♻️ 恢复到功能页",
            command=self._restore_history_record,
            width=130,
            height=32,
            bg_color=ModernStyle.PRIMARY,
            hover_color=ModernStyle.PRIMARY_DARK
        )
        # 初始隐藏恢复按钮
        
        self.history_dual_output = DualOutputFrame(
            right_panel,
            height=20,
            show_actions=False
        )
        self.history_dual_output.pack(fill=tk.BOTH, expand=True)
        
        paned.add(right_panel, minsize=400)
        
        self.current_history_record = None
        self._refresh_history()

    def _refresh_history(self):
        """刷新历史记录列表"""
        # 清空现有列表
        for item in self.history_tree.get_children():
            self.history_tree.delete(item)
            
        records = self.history.get_recent_records(limit=100)
        
        type_map = {
            "diagnose": "🔍 论文诊断",
            "optimize": "⚙️ 深度优化",
            "dedup": "📉 智能降重",
            "deai": "🤖 降AI痕迹",
            "deep_process": "⚡ 深度处理",
            "search": "🔎 学术搜索",
            "revision": "📝 退修助手"
        }
        
        for r in records:
            # 格式化时间
            try:
                dt = datetime.strptime(r['timestamp'], '%Y-%m-%d %H:%M:%S')
                time_str = dt.strftime('%m-%d %H:%M')
            except (ValueError, TypeError, KeyError):
                time_str = r.get('timestamp', 'N/A')
                
            action_name = type_map.get(r['action_type'], r['action_type'])
            preview = r['output_content'][:100].replace('\n', ' ')
            
            self.history_tree.insert("", tk.END, iid=str(r['id']), values=(time_str, action_name, preview))

    def _on_history_select(self, event):
        """选中历史记录"""
        selection = self.history_tree.selection()
        if not selection:
            return
            
        record_id = int(selection[0])
        # 从数据库获取完整记录
        r = self.history.get_record_by_id(record_id)
        
        if r:
            self.current_history_record = r
            
            # 显示详情
            self.history_dual_output.set_content(
                r['output_content'],
                r['report'] or "无分析报告"
            )
            
            # 显示恢复按钮
            self.restore_btn.pack(side=tk.RIGHT)
            
    def _clear_all_history(self):
        """清空所有历史记录"""
        if ConfirmDialog.show(self.root, "确认清空", "确定要永久删除所有历史记录吗？此操作不可撤销。"):
            if self.history.clear_history():
                self.notification.show("历史记录已清空", "success")
                self._refresh_history()
                self.history_dual_output.clear()
                self.restore_btn.pack_forget()

    def _restore_history_record(self):
        """将选中的历史记录恢复到对应的功能页面"""
        if not self.current_history_record:
            return
            
        r = self.current_history_record
        action_type = r['action_type']
        
        # 映射到页面 ID
        page_map = {
            "diagnose": "diagnose",
            "optimize": "optimize",
            "dedup": "dedup",
            "deai": "dedup",
            "deep_process": "dedup",
            "search": "search",
            "revision": "revision"
        }
        
        target_page = page_map.get(action_type)
        if not target_page:
            return
            
        # 填充内容
        if target_page == "diagnose":
            self.diag_input_comp.set_content(r['input_content'])
            self.diag_dual_output.set_content(r['output_content'], r['report'])
        elif target_page == "optimize":
            self.opt_input_comp.set_content(r['input_content'])
            # 如果有 context，恢复它
            if r.get('metadata') and 'context' in r['metadata']:
                self._toggle_opt_context(show=True)
                self.opt_context_input.set_content(r['metadata']['context'])
            self.opt_dual_output.set_content(r['output_content'], r['report'])
        elif target_page == "dedup":
            self.dedup_input_comp.set_content(r['input_content'])
            self.dedup_dual_output.set_content(r['output_content'], r['report'])
        elif target_page == "search":
            self.search_query.delete(0, tk.END)
            self.search_query.insert(0, r['input_content'])
            self.search_dual_output.set_content(r['output_content'], r['report'])
        elif target_page == "revision":
            self.rev_comments_comp.set_content(r['input_content'])
            self.rev_dual_output.set_content(r['output_content'], r['report'])
            
        # 切换页面
        self._show_page(target_page)
        self.notification.show(f"已恢复历史记录至「{self.status_bar.status_label.cget('text').split(': ')[1]}」", "success")

    def _load_ui_preferences(self):
        """加载 UI 偏好设置 (P2/P3)"""
        try:
            # 加载上次页面
            last_page = self.history.get_preference("last_page", "diagnose")
            
            # 加载深色模式 - 安全检查 dark_mode_var 是否已创建
            is_dark = self.history.get_preference("dark_mode", False)
            if hasattr(self, 'dark_mode_var'):
                self.dark_mode_var.set(is_dark)
            if is_dark:
                ModernStyle.set_dark_mode(True)
                ModernStyle.configure_styles(self.root)
            
            if last_page in self.pages:
                self._show_page(last_page)
        except Exception:
            pass

    def _save_ui_preference(self, key: str, value: Any):
        """保存 UI 偏好设置 (P2)"""
        try:
            self.history.set_preference(key, value)
        except Exception:
            pass

    def _on_dark_mode_toggle(self):
        """深色模式切换回调 (P3)"""
        is_dark = self.dark_mode_var.get()
        
        # 1. 更新全局样式常量
        ModernStyle.set_dark_mode(is_dark)
        ModernStyle.configure_styles(self.root)
        
        # 2. 持久化设置
        self._save_ui_preference("dark_mode", is_dark)
        
        # 3. 尝试更新主窗口背景
        self.root.configure(bg=ModernStyle.BG_MAIN)
        self.content_frame.configure(bg=ModernStyle.BG_MAIN)
        
        # 4. 提示用户
        self.notification.show(f"已切换至{'深色' if is_dark else '浅色'}模式，部分组件重启后效果更佳", "success")

    def _create_template_selector(self, parent, category, target_comp):
        """创建模板选择器组件 (P3)"""
        frame = tk.Frame(parent, bg=parent.cget("bg"))
        frame.pack(side=tk.LEFT, padx=15)
        
        tk.Label(
            frame,
            text="📋 模板:",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_XS),
            bg=parent.cget("bg"),
            fg=ModernStyle.TEXT_SECONDARY
        ).pack(side=tk.LEFT)
        
        var = tk.StringVar(value="选择预设...")
        
        def refresh_templates():
            templates = self.history.get_templates(category)
            names = [t['name'] for t in templates]
            combo['values'] = ["选择预设..."] + names
            return templates

        combo = ttk.Combobox(
            frame,
            textvariable=var,
            values=["选择预设..."],
            state="readonly",
            width=15,
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_XS)
        )
        combo.pack(side=tk.LEFT, padx=5)
        
        # 管理模板按钮
        manage_btn = tk.Label(
            frame,
            text="⚙️",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_XS),
            bg=parent.cget("bg"),
            fg=ModernStyle.TEXT_MUTED,
            cursor="hand2"
        )
        manage_btn.pack(side=tk.LEFT, padx=5)
        manage_btn.bind("<Button-1>", lambda e: self._manage_templates(category))
        Tooltip(manage_btn, "管理自定义模板")
        
        templates_data = refresh_templates()
        
        def on_select(event):
            name = var.get()
            if name == "选择预设...": return
            
            # 重新获取以确保最新
            curr_templates = self.history.get_templates(category)
            template = next((t for t in curr_templates if t['name'] == name), None)
            if template:
                current_val = target_comp.get_content()
                if current_val and not ConfirmDialog.show(self.root, "确认覆盖", "应用模板将覆盖当前输入内容，确定吗？"):
                    return
                target_comp.set_content(template['content'])
                self.notification.show(f"已应用模板: {name}", "info")
        
        combo.bind("<<ComboboxSelected>>", on_select)
        # 绑定点击事件以刷新列表
        combo.bind("<Button-1>", lambda e: refresh_templates())
        return combo

    def _save_as_template(self, category: str, source_comp):
        """将内容保存为模板 (P3)"""
        content = source_comp.get_content()
        if not content:
            self.notification.show("请先输入要保存的内容", "warning")
            return
            
        # 弹出简单对话框询问名称
        name_window = tk.Toplevel(self.root)
        name_window.title("保存模板")
        name_window.geometry("350x180")
        name_window.resizable(False, False)
        name_window.configure(bg=ModernStyle.BG_MAIN)
        name_window.transient(self.root)
        name_window.grab_set()
        
        content_frame = tk.Frame(name_window, bg=ModernStyle.BG_MAIN, padx=20, pady=20)
        content_frame.pack(fill=tk.BOTH, expand=True)
        
        tk.Label(content_frame, text="请输入模板名称:", bg=ModernStyle.BG_MAIN, font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_SM)).pack(anchor="w")
        name_entry = tk.Entry(content_frame, width=30)
        name_entry.pack(pady=10)
        name_entry.focus_set()
        
        def do_save():
            name = name_entry.get().strip()
            if not name:
                return
            if self.history.save_template(name, content, category):
                self.notification.show(f"模板「{name}」保存成功", "success")
                name_window.destroy()
            else:
                self.notification.show("保存模板失败", "error")
        
        ModernButton(content_frame, text="保存", command=do_save, width=100).pack(side=tk.LEFT, pady=10)
        ModernButton(content_frame, text="取消", command=name_window.destroy, width=100, bg_color=ModernStyle.BG_SECONDARY, text_color=ModernStyle.TEXT_PRIMARY).pack(side=tk.RIGHT, pady=10)

    def _manage_templates(self, category: str):
        """管理自定义模板对话框 (P3)"""
        manage_window = tk.Toplevel(self.root)
        manage_window.title("管理模板")
        manage_window.geometry("500x400")
        manage_window.configure(bg=ModernStyle.BG_MAIN)
        manage_window.transient(self.root)
        manage_window.grab_set()
        
        content = tk.Frame(manage_window, bg=ModernStyle.BG_MAIN, padx=20, pady=20)
        content.pack(fill=tk.BOTH, expand=True)
        
        tk.Label(content, text=f"管理模板 - {category}", font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_MD, "bold"), bg=ModernStyle.BG_MAIN).pack(anchor="w", pady=(0, 10))
        
        # 列表
        list_frame = tk.Frame(content, bg=ModernStyle.BG_MAIN)
        list_frame.pack(fill=tk.BOTH, expand=True)
        
        columns = ("name", "is_system")
        tree = ttk.Treeview(list_frame, columns=columns, show="headings", height=10)
        tree.heading("name", text="名称")
        tree.heading("is_system", text="类型")
        tree.column("name", width=300)
        tree.column("is_system", width=100)
        tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        scrollbar = ttk.Scrollbar(list_frame, orient="vertical", command=tree.yview)
        tree.configure(yscrollcommand=scrollbar.set)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        def load_list():
            for item in tree.get_children(): tree.delete(item)
            templates = self.history.get_templates(category)
            for t in templates:
                type_str = "系统" if t['is_system'] else "自定义"
                tree.insert("", tk.END, iid=str(t['id']), values=(t['name'], type_str))
                
        load_list()
        
        btn_frame = tk.Frame(content, bg=ModernStyle.BG_MAIN, pady=15)
        btn_frame.pack(fill=tk.X)
        
        def delete_selected():
            selection = tree.selection()
            if not selection: return
            tid = int(selection[0])
            # 检查是否为系统模板（前端再次确认）
            item = tree.item(selection[0])
            if item['values'][1] == "系统":
                self.notification.show("系统模板不可删除", "warning")
                return
                
            if ConfirmDialog.show(manage_window, "确认删除", "确定要删除该模板吗？"):
                if self.history.delete_template(tid):
                    self.notification.show("模板已删除", "success")
                    load_list()
                else:
                    self.notification.show("删除失败", "error")
        
        ModernButton(btn_frame, text="删除选中", command=delete_selected, bg_color=ModernStyle.ERROR, hover_color=ModernStyle.ERROR, width=120).pack(side=tk.LEFT)
        ModernButton(btn_frame, text="关闭", command=manage_window.destroy, width=100, bg_color=ModernStyle.BG_SECONDARY, text_color=ModernStyle.TEXT_PRIMARY).pack(side=tk.RIGHT)

    def _create_settings_page(self):
        """创建设置页面 - 优化版：分离API配置 + 模型拉取"""
        page = tk.Frame(self.content_frame, bg=ModernStyle.BG_MAIN)
        self.pages["settings"] = page
        
        self._create_page_header(page, "系统设置", "配置 AI 模型、API 密钥等参数")
        
        # 滚动区域
        canvas = tk.Canvas(page, bg=ModernStyle.BG_MAIN, highlightthickness=0)
        scrollbar = ttk.Scrollbar(page, orient="vertical", command=canvas.yview)
        scrollable_frame = tk.Frame(canvas, bg=ModernStyle.BG_MAIN)
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas_window = canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        
        def configure_canvas(event):
            canvas.itemconfig(canvas_window, width=event.width)
        canvas.bind('<Configure>', configure_canvas)
        
        canvas.configure(yscrollcommand=scrollbar.set)
        
        def _on_mousewheel(event):
            canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        canvas.bind_all("<MouseWheel>", _on_mousewheel)
        
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=ModernStyle.PADDING_XL)
        
        content = scrollable_frame
        
        # ============ 1. 语言模型配置 ============
        section1 = tk.Frame(content, bg=ModernStyle.BG_MAIN)
        section1.pack(fill=tk.X, pady=(0, 30))
        
        header1 = tk.Frame(section1, bg=ModernStyle.BG_MAIN)
        header1.pack(fill=tk.X, pady=(0, 15))
        
        tk.Label(
            header1,
            text="🤖 语言模型配置 (LLM)",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_LG, "bold"),
            bg=ModernStyle.BG_MAIN,
            fg=ModernStyle.TEXT_PRIMARY
        ).pack(side=tk.LEFT)
        
        ModernButton(
            header1,
            text="🔗 测试连接",
            command=self._test_llm_connection,
            width=120,
            height=36,
            bg_color=ModernStyle.BG_SECONDARY,
            hover_color=ModernStyle.BG_HOVER,
            text_color=ModernStyle.TEXT_PRIMARY,
            tooltip="验证 API 配置是否正确"
        ).pack(side=tk.RIGHT)
        
        llm_frame = tk.Frame(section1, bg=ModernStyle.BG_SECONDARY, padx=25, pady=25)
        llm_frame.pack(fill=tk.X)
        
        # 供应商选择
        row1 = tk.Frame(llm_frame, bg=ModernStyle.BG_SECONDARY)
        row1.pack(fill=tk.X, pady=10)
        
        tk.Label(
            row1,
            text="供应商:",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_MD),
            bg=ModernStyle.BG_SECONDARY,
            fg=ModernStyle.TEXT_PRIMARY,
            width=12,
            anchor="w"
        ).pack(side=tk.LEFT)
        
        self.llm_provider_var = tk.StringVar(value="OpenAI 兼容")
        providers = ["OpenAI 兼容", "DeepSeek", "硅基流动", "Ollama 本地", "自定义"]
        
        provider_combo = ttk.Combobox(
            row1,
            textvariable=self.llm_provider_var,
            values=providers,
            state="readonly",
            width=25,
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_SM)
        )
        provider_combo.pack(side=tk.LEFT, padx=12)
        provider_combo.bind("<<ComboboxSelected>>", self._on_llm_provider_change)
        
        tk.Label(
            row1,
            text="💡 切换供应商自动填充 API 地址",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_XS),
            bg=ModernStyle.BG_SECONDARY,
            fg=ModernStyle.TEXT_MUTED
        ).pack(side=tk.LEFT, padx=18)
        
        # API 地址
        row2 = tk.Frame(llm_frame, bg=ModernStyle.BG_SECONDARY)
        row2.pack(fill=tk.X, pady=10)
        
        tk.Label(
            row2,
            text="API 地址:",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_MD),
            bg=ModernStyle.BG_SECONDARY,
            fg=ModernStyle.TEXT_PRIMARY,
            width=12,
            anchor="w"
        ).pack(side=tk.LEFT)
        
        self.setting_llm_base = tk.Entry(
            row2,
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_MD),
            bg=ModernStyle.BG_MAIN,
            relief="flat",
            width=55
        )
        self.setting_llm_base.pack(side=tk.LEFT, padx=12, ipady=8)
        
        # API 密钥
        row3 = tk.Frame(llm_frame, bg=ModernStyle.BG_SECONDARY)
        row3.pack(fill=tk.X, pady=10)
        
        tk.Label(
            row3,
            text="API 密钥:",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_MD),
            bg=ModernStyle.BG_SECONDARY,
            fg=ModernStyle.TEXT_PRIMARY,
            width=12,
            anchor="w"
        ).pack(side=tk.LEFT)
        
        self.setting_llm_key = tk.Entry(
            row3,
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_MD),
            bg=ModernStyle.BG_MAIN,
            relief="flat",
            width=45,
            show="•"
        )
        self.setting_llm_key.pack(side=tk.LEFT, padx=12, ipady=8)
        
        self.show_llm_key = tk.BooleanVar(value=False)
        tk.Checkbutton(
            row3,
            text="显示",
            variable=self.show_llm_key,
            command=lambda: self.setting_llm_key.config(show="" if self.show_llm_key.get() else "•"),
            bg=ModernStyle.BG_SECONDARY,
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_SM)
        ).pack(side=tk.LEFT, padx=12)
        
        # 模型选择
        row4 = tk.Frame(llm_frame, bg=ModernStyle.BG_SECONDARY)
        row4.pack(fill=tk.X, pady=10)
        
        tk.Label(
            row4,
            text="模型名称:",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_MD),
            bg=ModernStyle.BG_SECONDARY,
            fg=ModernStyle.TEXT_PRIMARY,
            width=12,
            anchor="w"
        ).pack(side=tk.LEFT)
        
        self.setting_llm_model = ttk.Combobox(
            row4,
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_SM),
            width=35,
            values=["gpt-4o", "gpt-4o-mini", "gpt-4-turbo", "deepseek-chat", "deepseek-coder", 
                   "Qwen/Qwen2.5-72B-Instruct", "claude-3-5-sonnet-20241022"]
        )
        self.setting_llm_model.pack(side=tk.LEFT, padx=12)
        
        ModernButton(
            row4,
            text="📥 拉取模型列表",
            command=self._fetch_llm_models,
            width=140,
            height=32,
            bg_color=ModernStyle.BG_MAIN,
            hover_color=ModernStyle.BG_HOVER,
            text_color=ModernStyle.TEXT_PRIMARY,
            tooltip="从服务器获取可用模型"
        ).pack(side=tk.LEFT, padx=12)
        
        self.llm_status = tk.Label(
            row4,
            text="● 未配置",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_SM),
            bg=ModernStyle.BG_SECONDARY,
            fg=ModernStyle.WARNING
        )
        self.llm_status.pack(side=tk.LEFT, padx=12)
        
        # ============ 2. 嵌入模型配置 ============
        section2 = tk.Frame(content, bg=ModernStyle.BG_MAIN)
        section2.pack(fill=tk.X, pady=(0, 30))
        
        header2 = tk.Frame(section2, bg=ModernStyle.BG_MAIN)
        header2.pack(fill=tk.X, pady=(0, 15))
        
        tk.Label(
            header2,
            text="📊 嵌入模型配置 (Embedding)",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_LG, "bold"),
            bg=ModernStyle.BG_MAIN,
            fg=ModernStyle.TEXT_PRIMARY
        ).pack(side=tk.LEFT)
        
        # 嵌入模型启用开关 (可选功能)
        self.enable_embedding = tk.BooleanVar(value=False)
        
        self.use_same_api = tk.BooleanVar(value=True)
        tk.Checkbutton(
            header2,
            text="使用与语言模型相同的 API 配置",
            variable=self.use_same_api,
            command=self._toggle_embed_api,
            bg=ModernStyle.BG_MAIN,
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_SM)
        ).pack(side=tk.RIGHT)
        
        self.embed_frame = tk.Frame(section2, bg=ModernStyle.BG_SECONDARY, padx=25, pady=25)
        self.embed_frame.pack(fill=tk.X)
        
        # 嵌入模型 API 地址
        row_e1 = tk.Frame(self.embed_frame, bg=ModernStyle.BG_SECONDARY)
        row_e1.pack(fill=tk.X, pady=10)
        
        tk.Label(
            row_e1,
            text="API 地址:",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_MD),
            bg=ModernStyle.BG_SECONDARY,
            fg=ModernStyle.TEXT_PRIMARY,
            width=12,
            anchor="w"
        ).pack(side=tk.LEFT)
        
        self.setting_embed_base = tk.Entry(
            row_e1,
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_MD),
            bg=ModernStyle.BG_MAIN,
            relief="flat",
            width=55
        )
        self.setting_embed_base.pack(side=tk.LEFT, padx=12, ipady=8)
        
        # 嵌入模型 API 密钥
        row_e2 = tk.Frame(self.embed_frame, bg=ModernStyle.BG_SECONDARY)
        row_e2.pack(fill=tk.X, pady=10)
        
        tk.Label(
            row_e2,
            text="API 密钥:",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_MD),
            bg=ModernStyle.BG_SECONDARY,
            fg=ModernStyle.TEXT_PRIMARY,
            width=12,
            anchor="w"
        ).pack(side=tk.LEFT)
        
        self.setting_embed_key = tk.Entry(
            row_e2,
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_MD),
            bg=ModernStyle.BG_MAIN,
            relief="flat",
            width=45,
            show="•"
        )
        self.setting_embed_key.pack(side=tk.LEFT, padx=12, ipady=8)
        
        # 嵌入模型选择
        row_e3 = tk.Frame(self.embed_frame, bg=ModernStyle.BG_SECONDARY)
        row_e3.pack(fill=tk.X, pady=10)
        
        tk.Label(
            row_e3,
            text="模型名称:",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_MD),
            bg=ModernStyle.BG_SECONDARY,
            fg=ModernStyle.TEXT_PRIMARY,
            width=12,
            anchor="w"
        ).pack(side=tk.LEFT)
        
        self.setting_embed_model = ttk.Combobox(
            row_e3,
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_SM),
            width=35,
            values=["text-embedding-3-small", "text-embedding-3-large", "text-embedding-ada-002",
                   "BAAI/bge-m3", "BAAI/bge-large-zh-v1.5"]
        )
        self.setting_embed_model.pack(side=tk.LEFT, padx=12)
        
        ModernButton(
            row_e3,
            text="📥 拉取模型列表",
            command=self._fetch_embed_models,
            width=140,
            height=32,
            bg_color=ModernStyle.BG_MAIN,
            hover_color=ModernStyle.BG_HOVER,
            text_color=ModernStyle.TEXT_PRIMARY,
            tooltip="获取可用嵌入模型"
        ).pack(side=tk.LEFT, padx=12)
        
        # 初始状态：隐藏独立配置
        self._toggle_embed_api()
        
        # ============ 3. 数据存储配置 ============
        section3 = tk.Frame(content, bg=ModernStyle.BG_MAIN)
        section3.pack(fill=tk.X, pady=(0, 30))
        
        header3 = tk.Frame(section3, bg=ModernStyle.BG_MAIN)
        header3.pack(fill=tk.X, pady=(0, 15))
        
        tk.Label(
            header3,
            text="📁 数据存储配置",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_LG, "bold"),
            bg=ModernStyle.BG_MAIN,
            fg=ModernStyle.TEXT_PRIMARY
        ).pack(side=tk.LEFT)
        
        tk.Label(
            header3,
            text="💡 自定义存储位置可避免占用C盘空间",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_XS),
            bg=ModernStyle.BG_MAIN,
            fg=ModernStyle.TEXT_MUTED
        ).pack(side=tk.RIGHT)
        
        storage_frame = tk.Frame(section3, bg=ModernStyle.BG_SECONDARY, padx=25, pady=25)
        storage_frame.pack(fill=tk.X)
        
        # 数据目录
        row_s1 = tk.Frame(storage_frame, bg=ModernStyle.BG_SECONDARY)
        row_s1.pack(fill=tk.X, pady=10)
        
        tk.Label(
            row_s1,
            text="数据目录:",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_MD),
            bg=ModernStyle.BG_SECONDARY,
            fg=ModernStyle.TEXT_PRIMARY,
            width=12,
            anchor="w"
        ).pack(side=tk.LEFT)
        
        self.setting_data_dir = tk.Entry(
            row_s1,
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_MD),
            bg=ModernStyle.BG_MAIN,
            relief="flat",
            width=45
        )
        self.setting_data_dir.pack(side=tk.LEFT, padx=12, ipady=8)
        
        ModernButton(
            row_s1,
            text="📂 浏览",
            command=lambda: self._browse_directory("data_dir"),
            width=80,
            height=32,
            bg_color=ModernStyle.BG_MAIN,
            hover_color=ModernStyle.BG_HOVER,
            text_color=ModernStyle.TEXT_PRIMARY
        ).pack(side=tk.LEFT, padx=8)
        
        tk.Label(
            row_s1,
            text="(日志、缓存、向量库)",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_XS),
            bg=ModernStyle.BG_SECONDARY,
            fg=ModernStyle.TEXT_MUTED
        ).pack(side=tk.LEFT, padx=8)
        
        # 工作区目录
        row_s2 = tk.Frame(storage_frame, bg=ModernStyle.BG_SECONDARY)
        row_s2.pack(fill=tk.X, pady=10)
        
        tk.Label(
            row_s2,
            text="工作区目录:",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_MD),
            bg=ModernStyle.BG_SECONDARY,
            fg=ModernStyle.TEXT_PRIMARY,
            width=12,
            anchor="w"
        ).pack(side=tk.LEFT)
        
        self.setting_workspace_dir = tk.Entry(
            row_s2,
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_MD),
            bg=ModernStyle.BG_MAIN,
            relief="flat",
            width=45
        )
        self.setting_workspace_dir.pack(side=tk.LEFT, padx=12, ipady=8)
        
        ModernButton(
            row_s2,
            text="📂 浏览",
            command=lambda: self._browse_directory("workspace_dir"),
            width=80,
            height=32,
            bg_color=ModernStyle.BG_MAIN,
            hover_color=ModernStyle.BG_HOVER,
            text_color=ModernStyle.TEXT_PRIMARY
        ).pack(side=tk.LEFT, padx=8)
        
        tk.Label(
            row_s2,
            text="(导出文件存放位置)",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_XS),
            bg=ModernStyle.BG_SECONDARY,
            fg=ModernStyle.TEXT_MUTED
        ).pack(side=tk.LEFT, padx=8)
        
        # 当前存储位置显示
        row_s3 = tk.Frame(storage_frame, bg=ModernStyle.BG_SECONDARY)
        row_s3.pack(fill=tk.X, pady=(15, 5))
        
        self.storage_info_label = tk.Label(
            row_s3,
            text="",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_XS),
            bg=ModernStyle.BG_SECONDARY,
            fg=ModernStyle.TEXT_MUTED,
            wraplength=600,
            justify="left"
        )
        self.storage_info_label.pack(anchor="w")
        
        # 加载并显示当前存储位置
        self._update_storage_info()
        
        # ============ 4. 界面外观配置 (P3) ============
        section_ui = tk.Frame(content, bg=ModernStyle.BG_MAIN)
        section_ui.pack(fill=tk.X, pady=(0, 30))
        
        header_ui = tk.Frame(section_ui, bg=ModernStyle.BG_MAIN)
        header_ui.pack(fill=tk.X, pady=(0, 15))
        
        tk.Label(
            header_ui,
            text="🎨 界面外观配置",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_LG, "bold"),
            bg=ModernStyle.BG_MAIN,
            fg=ModernStyle.TEXT_PRIMARY
        ).pack(side=tk.LEFT)
        
        ui_frame = tk.Frame(section_ui, bg=ModernStyle.BG_SECONDARY, padx=25, pady=25)
        ui_frame.pack(fill=tk.X)
        
        row_ui1 = tk.Frame(ui_frame, bg=ModernStyle.BG_SECONDARY)
        row_ui1.pack(fill=tk.X, pady=10)
        
        tk.Label(
            row_ui1,
            text="深色模式:",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_MD),
            bg=ModernStyle.BG_SECONDARY,
            fg=ModernStyle.TEXT_PRIMARY,
            width=12,
            anchor="w"
        ).pack(side=tk.LEFT)
        
        self.dark_mode_var = tk.BooleanVar(value=False)
        tk.Checkbutton(
            row_ui1,
            text="开启深色主题",
            variable=self.dark_mode_var,
            command=self._on_dark_mode_toggle,
            bg=ModernStyle.BG_SECONDARY,
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_SM)
        ).pack(side=tk.LEFT, padx=12)

        # ============ 5. API 用量统计 (P2) ============
        section4 = tk.Frame(content, bg=ModernStyle.BG_MAIN)
        section4.pack(fill=tk.X, pady=(0, 30))
        
        header4 = tk.Frame(section4, bg=ModernStyle.BG_MAIN)
        header4.pack(fill=tk.X, pady=(0, 15))
        
        tk.Label(
            header4,
            text="📈 API 用量统计",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_LG, "bold"),
            bg=ModernStyle.BG_MAIN,
            fg=ModernStyle.TEXT_PRIMARY
        ).pack(side=tk.LEFT)
        
        ModernButton(
            header4,
            text="🔄 刷新统计",
            command=self._refresh_usage_stats,
            width=100,
            height=32,
            bg_color=ModernStyle.BG_SECONDARY,
            hover_color=ModernStyle.BG_HOVER,
            text_color=ModernStyle.TEXT_PRIMARY
        ).pack(side=tk.RIGHT)
        
        self.usage_frame = tk.Frame(section4, bg=ModernStyle.BG_SECONDARY, padx=25, pady=25)
        self.usage_frame.pack(fill=tk.X)
        
        self.usage_label = tk.Label(
            self.usage_frame,
            text="正在加载统计信息...",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_MD),
            bg=ModernStyle.BG_SECONDARY,
            fg=ModernStyle.TEXT_PRIMARY,
            justify="left"
        )
        self.usage_label.pack(anchor="w")
        
        # ============ 5. 保存按钮 ============
        btn_frame = tk.Frame(content, bg=ModernStyle.BG_MAIN)
        btn_frame.pack(fill=tk.X, pady=30)
        
        ModernButton(
            btn_frame,
            text="💾 保存配置",
            command=self._save_settings,
            width=160,
            height=48
        ).pack(side=tk.LEFT)
        
        ModernButton(
            btn_frame,
            text="恢复默认",
            command=self._reset_settings,
            width=120,
            height=48,
            bg_color=ModernStyle.BG_SECONDARY,
            hover_color=ModernStyle.BG_HOVER,
            text_color=ModernStyle.TEXT_PRIMARY
        ).pack(side=tk.LEFT, padx=18)
        
        # 加载现有设置
        self._load_settings()
    
    def _toggle_embed_api(self):
        """切换嵌入模型配置显示"""
        # 保存当前输入值（如果存在）
        saved_embed_base = ""
        saved_embed_key = ""
        saved_embed_model = ""
        
        if hasattr(self, 'setting_embed_base') and hasattr(self.setting_embed_base, 'winfo_exists'):
            try:
                if self.setting_embed_base.winfo_exists():
                    saved_embed_base = self.setting_embed_base.get()
            except Exception:
                pass
        if hasattr(self, 'setting_embed_key') and hasattr(self.setting_embed_key, 'winfo_exists'):
            try:
                if self.setting_embed_key.winfo_exists():
                    saved_embed_key = self.setting_embed_key.get()
            except Exception:
                pass
        if hasattr(self, 'setting_embed_model'):
            try:
                saved_embed_model = self.setting_embed_model.get()
            except Exception:
                pass
        
        # 清除现有内容
        for widget in self.embed_frame.winfo_children():
            if isinstance(widget, tk.Frame):
                widget.destroy()
        
        # 嵌入模型为可选功能，默认使用语言模型的 API
        if self.use_same_api.get():
            # 使用相同API - 只显示模型选择
            row_e3 = tk.Frame(self.embed_frame, bg=ModernStyle.BG_SECONDARY)
            row_e3.pack(fill=tk.X, pady=10)
            
            tk.Label(
                row_e3,
                text="模型名称:",
                font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_MD),
                bg=ModernStyle.BG_SECONDARY,
                fg=ModernStyle.TEXT_PRIMARY,
                width=12,
                anchor="w"
            ).pack(side=tk.LEFT)
            
            self.setting_embed_model = ttk.Combobox(
                row_e3,
                font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_SM),
                width=35,
                values=["text-embedding-3-small", "text-embedding-3-large", "BAAI/bge-m3"]
            )
            self.setting_embed_model.pack(side=tk.LEFT, padx=12)
            
            # 恢复保存的模型值
            if saved_embed_model:
                self.setting_embed_model.set(saved_embed_model)
            
            tk.Label(
                row_e3,
                text="💡 将使用语言模型的 API 地址和密钥",
                font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_XS),
                bg=ModernStyle.BG_SECONDARY,
                fg=ModernStyle.TEXT_MUTED
            ).pack(side=tk.LEFT, padx=18)
        else:
            # 使用独立API - 显示完整配置
            # API 地址
            row_e1 = tk.Frame(self.embed_frame, bg=ModernStyle.BG_SECONDARY)
            row_e1.pack(fill=tk.X, pady=10)
            
            tk.Label(
                row_e1,
                text="API 地址:",
                font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_MD),
                bg=ModernStyle.BG_SECONDARY,
                fg=ModernStyle.TEXT_PRIMARY,
                width=12,
                anchor="w"
            ).pack(side=tk.LEFT)
            
            self.setting_embed_base = tk.Entry(
                row_e1,
                font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_MD),
                bg=ModernStyle.BG_MAIN,
                relief="flat",
                width=55
            )
            self.setting_embed_base.pack(side=tk.LEFT, padx=12, ipady=8)
            
            # 恢复保存的值
            if saved_embed_base:
                self.setting_embed_base.insert(0, saved_embed_base)
            
            # API 密钥
            row_e2 = tk.Frame(self.embed_frame, bg=ModernStyle.BG_SECONDARY)
            row_e2.pack(fill=tk.X, pady=10)
            
            tk.Label(
                row_e2,
                text="API 密钥:",
                font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_MD),
                bg=ModernStyle.BG_SECONDARY,
                fg=ModernStyle.TEXT_PRIMARY,
                width=12,
                anchor="w"
            ).pack(side=tk.LEFT)
            
            self.setting_embed_key = tk.Entry(
                row_e2,
                font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_MD),
                bg=ModernStyle.BG_MAIN,
                relief="flat",
                width=45,
                show="•"
            )
            self.setting_embed_key.pack(side=tk.LEFT, padx=12, ipady=8)
            
            # 恢复保存的值
            if saved_embed_key:
                self.setting_embed_key.insert(0, saved_embed_key)
            
            # 模型选择
            row_e3 = tk.Frame(self.embed_frame, bg=ModernStyle.BG_SECONDARY)
            row_e3.pack(fill=tk.X, pady=10)
            
            tk.Label(
                row_e3,
                text="模型名称:",
                font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_MD),
                bg=ModernStyle.BG_SECONDARY,
                fg=ModernStyle.TEXT_PRIMARY,
                width=12,
                anchor="w"
            ).pack(side=tk.LEFT)
            
            self.setting_embed_model = ttk.Combobox(
                row_e3,
                font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_SM),
                width=35,
                values=["text-embedding-3-small", "text-embedding-3-large", "BAAI/bge-m3"]
            )
            self.setting_embed_model.pack(side=tk.LEFT, padx=12)
            
            # 恢复保存的模型值
            if saved_embed_model:
                self.setting_embed_model.set(saved_embed_model)
    
    def _on_llm_provider_change(self, event=None):
        """切换供应商时自动填充"""
        provider = self.llm_provider_var.get()
        
        presets = {
            "OpenAI 兼容": ("https://api.openai.com/v1", "gpt-4o-mini", "text-embedding-3-small"),
            "DeepSeek": ("https://api.deepseek.com/v1", "deepseek-chat", "text-embedding-3-small"),
            "硅基流动": ("https://api.siliconflow.cn/v1", "Qwen/Qwen2.5-72B-Instruct", "BAAI/bge-m3"),
            "Ollama 本地": ("http://localhost:11434/v1", "llama3.2", "nomic-embed-text"),
            "自定义": ("", "", ""),
        }
        
        if provider in presets:
            base, model, embed = presets[provider]
            self.setting_llm_base.delete(0, tk.END)
            self.setting_llm_base.insert(0, base)
            self.setting_llm_model.set(model)
            self.setting_embed_model.set(embed)
    
    def _fetch_llm_models(self):
        """拉取语言模型列表"""
        api_base = self.setting_llm_base.get().strip()
        api_key = self.setting_llm_key.get().strip()
        
        if not api_base or not api_key:
            self.notification.show("请先填写 API 地址和密钥", "warning")
            return
        
        def do_fetch(check_cancel):
            if OpenAI is None:
                raise ImportError("未安装 openai 库")
            client = OpenAI(base_url=api_base, api_key=api_key)
            models = client.models.list()
            
            model_ids = [m.id for m in models.data]
            model_ids.sort()
            return model_ids
        
        def on_complete(model_ids):
            self.setting_llm_model.config(values=model_ids)
            self.notification.show(f"成功获取到 {len(model_ids)} 个模型", "success")
            
        self.task_manager.submit(do_fetch, on_complete=on_complete, task_name="fetch_models")
    
    def _fetch_embed_models(self):
        """拉取嵌入模型列表"""
        if self.use_same_api.get():
            api_base = self.setting_llm_base.get().strip()
            api_key = self.setting_llm_key.get().strip()
        else:
            api_base = self.setting_embed_base.get().strip()
            api_key = self.setting_embed_key.get().strip()
        
        if not api_base or not api_key:
            self.notification.show("请先填写 API 地址和密钥", "warning")
            return
        
        def do_fetch(check_cancel):
            if OpenAI is None:
                raise ImportError("未安装 openai 库")
            client = OpenAI(base_url=api_base, api_key=api_key)
            models = client.models.list()
            embed_ids = [m.id for m in models.data if 'embed' in m.id.lower() or 'bge' in m.id.lower()]
            embed_ids.sort()
            return embed_ids
        
        def on_complete(embed_ids):
            if embed_ids:
                self.setting_embed_model.config(values=embed_ids)
                self.notification.show(f"成功获取到 {len(embed_ids)} 个嵌入模型", "success")
            else:
                self.notification.show("未找到嵌入模型，请手动输入", "warning")
            
        self.task_manager.submit(do_fetch, on_complete=on_complete, task_name="fetch_embed_models")
    
    def _test_llm_connection(self):
        """测试语言模型连接"""
        api_base = self.setting_llm_base.get().strip()
        api_key = self.setting_llm_key.get().strip()
        model = self.setting_llm_model.get().strip()
        
        if not api_base or not api_key:
            self.notification.show("请先填写 API 地址和密钥", "warning")
            return
        
        def do_test(check_cancel):
            if OpenAI is None:
                raise ImportError("未安装 openai 库")
            client = OpenAI(base_url=api_base, api_key=api_key)
            client.chat.completions.create(
                model=model,
                messages=[{"role": "user", "content": "Hi"}],
                max_tokens=5
            )
            return True
        
        def on_complete(res):
            self.llm_status.config(text="● 已连接", fg=ModernStyle.SUCCESS)
            self.notification.show("连接成功！API 配置有效。", "success")
            
        def on_error(err):
            self.llm_status.config(text="● 连接失败", fg=ModernStyle.ERROR)
            self.notification.show(f"连接失败: {str(err)}", "error")
            
        self.task_manager.submit(do_test, on_complete=on_complete, on_error=on_error, task_name="test_connection")
    
    def _browse_directory(self, target: str):
        """浏览选择目录"""
        directory = filedialog.askdirectory(title="选择目录")
        if directory:
            if target == "data_dir":
                self.setting_data_dir.delete(0, tk.END)
                self.setting_data_dir.insert(0, directory)
            elif target == "workspace_dir":
                self.setting_workspace_dir.delete(0, tk.END)
                self.setting_workspace_dir.insert(0, directory)
    
    def _update_storage_info(self):
        """更新存储位置信息显示"""
        try:
            from config.settings import settings
            info_text = f"📍 当前数据目录: {settings.data_dir}\n📍 当前工作区: {settings.workspace_dir}"
            self.storage_info_label.config(text=info_text)
            
            # 填充输入框
            self.setting_data_dir.delete(0, tk.END)
            self.setting_data_dir.insert(0, settings.data_dir)
            self.setting_workspace_dir.delete(0, tk.END)
            self.setting_workspace_dir.insert(0, settings.workspace_dir)
        except Exception:
            pass
    
    def _refresh_usage_stats(self):
        """刷新 API 用量统计 (P2)"""
        stats = self.history.get_usage_summary()
        text = (
            f"● 总请求次数: {stats['total_requests']} 次\n"
            f"● 总消耗 Token: {stats['total_tokens']:,}\n"
            f"● 预估总成本: ${stats['total_cost']:.4f} (基于标准费率)"
        )
        self.usage_label.config(text=text)

    def _reset_settings(self):
        """重置设置"""
        if ConfirmDialog.show(self.root, "确认重置", "确定要重置所有设置吗？"):
            self.setting_llm_base.delete(0, tk.END)
            self.setting_llm_key.delete(0, tk.END)
            self.setting_llm_model.set("")
            # 安全访问嵌入模型控件
            if hasattr(self, 'setting_embed_base') and hasattr(self.setting_embed_base, 'winfo_exists') and self.setting_embed_base.winfo_exists():
                self.setting_embed_base.delete(0, tk.END)
            if hasattr(self, 'setting_embed_key') and hasattr(self.setting_embed_key, 'winfo_exists') and self.setting_embed_key.winfo_exists():
                self.setting_embed_key.delete(0, tk.END)
            if hasattr(self, 'setting_embed_model'):
                self.setting_embed_model.set("")
            # 重置存储目录
            if hasattr(self, 'setting_data_dir'):
                self.setting_data_dir.delete(0, tk.END)
            if hasattr(self, 'setting_workspace_dir'):
                self.setting_workspace_dir.delete(0, tk.END)
            self.llm_provider_var.set("OpenAI 兼容")
            self.llm_status.config(text="● 未配置", fg=ModernStyle.WARNING)
    
    def _check_first_run(self):
        """首次运行检查 - 引导用户配置API"""
        try:
            from config.settings import settings
            if not settings.llm_api_key or not settings.llm_api_base:
                self.api_configured = False
                self._show_first_run_guide()
            else:
                self.api_configured = True
        except Exception:
            self.api_configured = False
            self._show_first_run_guide()
    
    def _show_first_run_guide(self):
        """显示首次使用引导"""
        guide_window = tk.Toplevel(self.root)
        guide_window.title("🎉 欢迎使用 EconPaper Pro")
        guide_window.geometry("550x450")
        guide_window.resizable(False, False)
        guide_window.configure(bg=ModernStyle.BG_MAIN)
        guide_window.transient(self.root)
        guide_window.grab_set()
        
        # 居中显示
        guide_window.update_idletasks()
        x = (guide_window.winfo_screenwidth() - 550) // 2
        y = (guide_window.winfo_screenheight() - 450) // 2
        guide_window.geometry(f"+{x}+{y}")
        
        # 内容
        content = tk.Frame(guide_window, bg=ModernStyle.BG_MAIN, padx=40, pady=30)
        content.pack(fill=tk.BOTH, expand=True)
        
        tk.Label(
            content,
            text="🎉 欢迎使用 EconPaper Pro!",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_XL, "bold"),
            bg=ModernStyle.BG_MAIN,
            fg=ModernStyle.TEXT_PRIMARY
        ).pack(pady=(0, 20))
        
        tk.Label(
            content,
            text="检测到您还未配置 AI 模型，请先完成以下设置：",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_MD),
            bg=ModernStyle.BG_MAIN,
            fg=ModernStyle.TEXT_SECONDARY
        ).pack(pady=(0, 25))
        
        # 步骤说明
        steps = [
            ("1️⃣", "准备 API 密钥", "从 OpenAI、DeepSeek 或硅基流动等平台获取"),
            ("2️⃣", "进入设置页面", "点击左侧「系统设置」"),
            ("3️⃣", "填写配置", "选择供应商 → 填写 API 密钥 → 保存"),
            ("4️⃣", "开始使用", "配置完成后即可使用所有功能"),
        ]
        
        for icon, title, desc in steps:
            step_frame = tk.Frame(content, bg=ModernStyle.BG_SECONDARY, padx=15, pady=12)
            step_frame.pack(fill=tk.X, pady=5)
            
            tk.Label(
                step_frame,
                text=icon,
                font=(ModernStyle.FONT_FAMILY, 16),
                bg=ModernStyle.BG_SECONDARY
            ).pack(side=tk.LEFT, padx=(0, 12))
            
            text_frame = tk.Frame(step_frame, bg=ModernStyle.BG_SECONDARY)
            text_frame.pack(side=tk.LEFT, fill=tk.X, expand=True)
            
            tk.Label(
                text_frame,
                text=title,
                font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_MD, "bold"),
                bg=ModernStyle.BG_SECONDARY,
                fg=ModernStyle.TEXT_PRIMARY,
                anchor="w"
            ).pack(anchor="w")
            
            tk.Label(
                text_frame,
                text=desc,
                font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_XS),
                bg=ModernStyle.BG_SECONDARY,
                fg=ModernStyle.TEXT_MUTED,
                anchor="w"
            ).pack(anchor="w")
        
        # 按钮
        btn_frame = tk.Frame(content, bg=ModernStyle.BG_MAIN)
        btn_frame.pack(fill=tk.X, pady=(25, 0))
        
        def go_to_settings():
            guide_window.destroy()
            self._show_page("settings")
        
        ModernButton(
            btn_frame,
            text="前往设置",
            command=go_to_settings,
            width=150,
            height=45
        ).pack(side=tk.LEFT)
        
        ModernButton(
            btn_frame,
            text="稍后配置",
            command=guide_window.destroy,
            width=120,
            height=45,
            bg_color=ModernStyle.BG_SECONDARY,
            hover_color=ModernStyle.BG_HOVER,
            text_color=ModernStyle.TEXT_PRIMARY
        ).pack(side=tk.LEFT, padx=15)
    
    def _show_about_dialog(self):
        """显示关于对话框"""
        about_window = tk.Toplevel(self.root)
        about_window.title("关于 EconPaper Pro")
        about_window.geometry("450x380")
        about_window.resizable(False, False)
        about_window.configure(bg=ModernStyle.BG_MAIN)
        about_window.transient(self.root)
        about_window.grab_set()
        
        # 居中显示
        about_window.update_idletasks()
        x = (about_window.winfo_screenwidth() - 450) // 2
        y = (about_window.winfo_screenheight() - 380) // 2
        about_window.geometry(f"+{x}+{y}")
        
        content = tk.Frame(about_window, bg=ModernStyle.BG_MAIN, padx=40, pady=30)
        content.pack(fill=tk.BOTH, expand=True)
        
        # Logo
        tk.Label(
            content,
            text="📚",
            font=(ModernStyle.FONT_FAMILY, 48),
            bg=ModernStyle.BG_MAIN
        ).pack()
        
        tk.Label(
            content,
            text="EconPaper Pro",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_XL, "bold"),
            bg=ModernStyle.BG_MAIN,
            fg=ModernStyle.TEXT_PRIMARY
        ).pack(pady=(10, 5))
        
        tk.Label(
            content,
            text=f"版本 {VERSION}",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_MD),
            bg=ModernStyle.BG_MAIN,
            fg=ModernStyle.TEXT_SECONDARY
        ).pack()
        
        tk.Label(
            content,
            text="经管学术论文智能助手",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_SM),
            bg=ModernStyle.BG_MAIN,
            fg=ModernStyle.TEXT_MUTED
        ).pack(pady=(15, 20))
        
        # 功能列表
        features = "✅ 论文诊断  ✅ 深度优化  ✅ 降重降AI\n✅ 学术搜索  ✅ 退修助手  ✅ 期刊过滤"
        tk.Label(
            content,
            text=features,
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_SM),
            bg=ModernStyle.BG_MAIN,
            fg=ModernStyle.TEXT_PRIMARY,
            justify="center"
        ).pack(pady=(0, 20))
        
        # 链接
        link_frame = tk.Frame(content, bg=ModernStyle.BG_MAIN)
        link_frame.pack()
        
        tk.Label(
            link_frame,
            text="📖 使用帮助",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_SM),
            bg=ModernStyle.BG_MAIN,
            fg=ModernStyle.PRIMARY,
            cursor="hand2"
        ).pack(side=tk.LEFT, padx=15)
        
        tk.Label(
            link_frame,
            text="🐛 反馈问题",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_SM),
            bg=ModernStyle.BG_MAIN,
            fg=ModernStyle.PRIMARY,
            cursor="hand2"
        ).pack(side=tk.LEFT, padx=15)
        
        # 关闭按钮
        ModernButton(
            content,
            text="关闭",
            command=about_window.destroy,
            width=100,
            height=40,
            bg_color=ModernStyle.BG_SECONDARY,
            hover_color=ModernStyle.BG_HOVER,
            text_color=ModernStyle.TEXT_PRIMARY
        ).pack(pady=(20, 0))
    
    def _check_api_before_action(self, action_name: str) -> bool:
        """执行操作前检查 API 配置"""
        if not self.api_configured:
            try:
                from config.settings import settings
                if settings.llm_api_key and settings.llm_api_base:
                    self.api_configured = True
                    return True
            except Exception:
                pass
            
            result = ConfirmDialog.show(
                self.root,
                "需要配置 API",
                f"使用「{action_name}」功能需要先配置 AI 模型。\n\n是否现在前往设置？",
                confirm_text="前往设置",
                cancel_text="稍后再说"
            )
            if result:
                self._show_page("settings")
            return False
        return True
    
    # ==================== 核心功能方法 ====================
    
    def _export_result(self, content: str, default_name: str):
        """导出结果到文件"""
        if not content or not content.strip():
            self.notification.show("没有可导出的内容", "warning")
            return
        
        file_types = [
            ("文本文件", "*.txt"),
            ("Markdown", "*.md"),
            ("Word 文档", "*.docx"),
            ("所有文件", "*.*")
        ]
        
        file_path = filedialog.asksaveasfilename(
            title="导出结果",
            defaultextension=".txt",
            filetypes=file_types,
            initialfile=f"{default_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        )
        
        if file_path:
            try:
                if file_path.lower().endswith(".docx"):
                    self._export_as_docx(content, file_path, default_name)
                else:
                    with open(file_path, "w", encoding="utf-8") as f:
                        f.write(content)
                    self.notification.show(f"已导出到: {os.path.basename(file_path)}", "success")
            except Exception as e:
                self.notification.show(f"导出失败: {e}", "error")

    def _export_as_docx(self, content: str, file_path: str, title: str):
        """将内容导出为专业 Word 文档 (P3)"""
        if not HAS_DOCX:
            self.notification.show("未安装 python-docx 库，无法导出 Word 格式", "error")
            return
            
        doc = docx.Document()
        doc.add_heading(title, 0)
        
        # 简单转换：按行处理，支持基础 Markdown 标题识别
        for line in content.split('\n'):
            line = line.strip()
            if not line:
                continue
            
            if line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            else:
                doc.add_paragraph(line)
        
        doc.save(file_path)
        self.notification.show(f"Word 文档已保存: {os.path.basename(file_path)}", "success")
    
    def _copy_to_clipboard(self, content: str):
        """复制内容到剪贴板"""
        if not content or not content.strip():
            self.notification.show("没有可复制的内容", "warning")
            return
        
        self.root.clipboard_clear()
        self.root.clipboard_append(content.strip())
        self.notification.show("已复制到剪贴板", "success")
    
    def _select_file(self, target: str):
        """选择文件 - 支持多选 (P3)"""
        file_paths = filedialog.askopenfilenames(
            title="选择论文文件 (支持多选)",
            filetypes=[
                ("支持的格式", "*.pdf;*.docx"),
                ("PDF 文件", "*.pdf"),
                ("Word 文档", "*.docx"),
                ("所有文件", "*.*")
            ]
        )
        
        if file_paths:
            count = len(file_paths)
            first_name = os.path.basename(file_paths[0])
            display_text = f"✓ {first_name}" + (f" 等 {count} 个文件" if count > 1 else "")
            
            if target == "diagnose":
                self.diag_file_paths = list(file_paths)
                self.diag_file_path = file_paths[0] # 保持兼容性
                self.diag_file_label.config(text=display_text, fg=ModernStyle.SUCCESS)
            elif target == "optimize":
                self.opt_file_paths = list(file_paths)
                self.opt_file_path = file_paths[0] # 保持兼容性
                self.opt_file_label.config(text=display_text, fg=ModernStyle.SUCCESS)
            elif target == "dedup":
                self.dedup_file_paths = list(file_paths)
                self.dedup_file_label.config(text=display_text, fg=ModernStyle.SUCCESS)
    
    def _set_result(self, widget: scrolledtext.ScrolledText, text: str):
        """设置结果文本"""
        def update():
            widget.config(state=tk.NORMAL)
            widget.delete("1.0", tk.END)
            widget.insert("1.0", text)
            widget.config(state=tk.DISABLED)
        self._safe_update(update)
    
    def _run_diagnose(self):
        """运行诊断 - 支持批量处理 (P3)"""
        if not self._check_api_before_action("论文诊断"):
            return
        
        # 确定待处理文件列表
        files_to_process = self.diag_file_paths if self.diag_file_paths else []
        is_batch = len(files_to_process) > 1
        
        if not files_to_process:
            text = self.diag_text.get("1.0", tk.END).strip()
            if not text:
                self.notification.show("请上传文件或粘贴论文内容", "warning")
                return
            # 文本模式，伪造一个列表以统一流程
            process_queue = [(None, text, None)]
        else:
            process_queue = []
            for fp in files_to_process:
                f_type = "pdf" if fp.lower().endswith(".pdf") else "docx"
                process_queue.append((fp, None, f_type))

        self.diag_dual_output.clear()
        
        def do_batch_diagnose(check_cancel):
            from agents.master import MasterAgent
            from agents.diagnostic import DiagnosticAgent
            agent = MasterAgent()
            diagnostic = DiagnosticAgent()
            
            total_files = len(process_queue)
            self._safe_update(lambda: self.precise_progress["diagnose"].start(total_files, "准备开始诊断..."))
            
            batch_results = []
            
            for i, (f_path, raw_text, f_type) in enumerate(process_queue, 1):
                if check_cancel(): return None
                
                fname = os.path.basename(f_path) if f_path else "粘贴的文本"
                # 修复 Lambda 闭包陷阱：使用默认参数捕获当前值
                self._safe_update(lambda idx=i, name=fname, total=total_files: self.precise_progress["diagnose"].update(idx, f"正在处理 ({idx}/{total}): {name}"))
                
                try:
                    content = raw_text
                    if f_path:
                        with open(f_path, "rb") as f:
                            content = f.read()
                    
                    # 诊断单个文件
                    report = agent.diagnose_only(content, file_type=f_type)
                    formatted = diagnostic.format_report(report)
                    
                    res_obj = {
                        'filename': fname,
                        'content': content if isinstance(content, str) else f"（{fname} 文件内容已解析）",
                        'report': f"### 文件: {fname}\n📊 评分: {report.overall_score:.1f}/10\n\n{formatted}"
                    }
                    batch_results.append(res_obj)
                    
                    # 保存每条历史记录
                    self.history.save_record(
                        action_type="diagnose",
                        input_content=res_obj['content'],
                        output_content=res_obj['content'],
                        report=res_obj['report'],
                        metadata={'file_path': f_path}
                    )
                except Exception as e:
                    batch_results.append({'filename': fname, 'content': '', 'report': f"### 文件: {fname}\n❌ 诊断失败: {e}"})

            return batch_results
        
        def on_complete(results):
            if results:
                # 汇总结果显示
                all_reports = []
                for r in results:
                    all_reports.append(r['report'])
                
                summary = "\n\n" + "="*50 + "\n\n".join(all_reports)
                
                if is_batch:
                    self.diag_dual_output.set_content(
                        f"已完成 {len(results)} 个文件的批量诊断。详细报告请查看「分析报告」选项卡。",
                        f"# 批量诊断汇总报告\n\n" + summary
                    )
                else:
                    self.diag_dual_output.set_result(results[0])
                
                self.notification.show(f"诊断完成 (共 {len(results)} 项)", "success")
                self.status_bar.set_status("诊断完成", "success")
            self.precise_progress["diagnose"].stop()
            self.is_processing = False

        def on_error(err):
            self.notification.show(f"处理出错: {err}", "error")
            self.precise_progress["diagnose"].stop(success=False)
            self.is_processing = False
            
        self.is_processing = True
        self.status_bar.set_status("正在处理任务...", "warning")
        self.task_manager.submit(do_batch_diagnose, on_complete=on_complete, on_error=on_error, task_name="diagnose")
    
    def _run_optimize(self):
        """运行优化 - 支持批量处理 (P3)"""
        if not self._check_api_before_action("深度优化"):
            return
        
        files_to_process = self.opt_file_paths if self.opt_file_paths else []
        is_batch = len(files_to_process) > 1
        
        if not files_to_process:
            text = self.opt_input_comp.get_content()
            if not text:
                self.notification.show("请上传文件或粘贴论文内容", "warning")
                return
            
            # 只有单个文本时，开启交互式流式优化
            sections = [k for k, v in self.opt_sections.items() if v.get()]
            if not sections:
                self.notification.show("请至少选择一个要优化的章节", "warning")
                return
                
            self._run_optimize_stream(text, sections)
            return
            
        # 批量文件处理流程
        process_queue = [(fp, None, "pdf" if fp.lower().endswith(".pdf") else "docx") for fp in files_to_process]
        
        sections = [k for k, v in self.opt_sections.items() if v.get()]
        if not sections:
            self.notification.show("请至少选择一个要优化的章节", "warning")
            return
        
        stage = self.opt_stage.get()
        journal = self.opt_journal.get() or None
        self.opt_dual_output.clear()
        
        def do_batch_optimize(check_cancel):
            from agents.optimizer import OptimizerAgent
            from parsers.structure import StructureRecognizer
            recognizer = StructureRecognizer()
            optimizer = OptimizerAgent(stage=stage)
            
            total_files = len(process_queue)
            self._safe_update(lambda: self.precise_progress["optimize"].start(total_files, "开始批量优化..."))
            
            batch_results = []
            for idx, (f_path, raw_text, f_type) in enumerate(process_queue, 1):
                if check_cancel(): return None
                fname = os.path.basename(f_path) if f_path else "粘贴的文本"
                # 修复 Lambda 闭包陷阱：使用默认参数捕获当前值
                self._safe_update(lambda i=idx, name=fname, total=total_files: self.precise_progress["optimize"].update(i, f"优化中 ({i}/{total}): {name}"))
                
                try:
                    content = raw_text
                    if f_path:
                        with open(f_path, "rb") as f: content = f.read()
                    
                    paper_structure = recognizer.recognize(content if isinstance(content, str) else f"（{fname} 内容已解析）")
                    content_parts = []
                    
                    for s_idx, section in enumerate(sections):
                        if check_cancel(): return None
                        s_content = paper_structure.get(section, content if len(sections)==1 else "")
                        
                        # 批量模式不进行流式渲染以保证性能，仅最后汇总
                        res_text = optimizer.optimize_single_section(section, s_content).optimized
                        content_parts.append(f"## {section.upper()}\n\n{res_text}")
                    
                    final_content = "\n\n".join(content_parts)
                    res_obj = {
                        'filename': fname,
                        'content': final_content,
                        'report': f"### 文件: {fname}\n阶段: {stage}\n章节: {', '.join(sections)}"
                    }
                    batch_results.append(res_obj)
                    
                    # 保存历史
                    self.history.save_record(
                        action_type="optimize",
                        input_content=content if isinstance(content, str) else f"File: {fname}",
                        output_content=final_content,
                        report=res_obj['report'],
                        metadata={'stage': stage, 'sections': sections}
                    )
                except Exception as e:
                    batch_results.append({'filename': fname, 'content': '', 'report': f"### 文件: {fname}\n❌ 优化失败: {e}"})
            
            return batch_results

        def on_complete(results):
            if results:
                if is_batch:
                    all_content = []
                    all_reports = []
                    for r in results:
                        all_content.append(f"--- 文件: {r['filename']} ---\n{r['content']}")
                        all_reports.append(r['report'])
                    
                    self.opt_dual_output.set_content(
                        "\n\n".join(all_content),
                        "# 批量优化报告\n\n" + "\n\n".join(all_reports)
                    )
                else:
                    self.opt_dual_output.set_result(results[0])
                
                self.notification.show(f"优化完成 (共 {len(results)} 项)", "success")
                self.status_bar.set_status("优化完成", "success")
            self.precise_progress["optimize"].stop()
            self.is_processing = False

        def on_error(err):
            self.notification.show(f"优化任务出错: {err}", "error")
            self.precise_progress["optimize"].stop(success=False)
            self.is_processing = False
            
        self.is_processing = True
        self.status_bar.set_status("正在处理...", "warning")
        self.task_manager.submit(do_batch_optimize, on_complete=on_complete, on_error=on_error, task_name="optimize")

    def _run_optimize_stream(self, text: str, sections: List[str]):
        """流式运行优化 (针对单个文本输入)"""
        from agents.optimizer import OptimizerAgent
        stage = self.opt_stage.get()
        journal = self.opt_journal.get() or "通用"
        
        # 获取背景参考
        context = self.opt_context_input.get_content() if self.opt_context_visible.get() else ""
        full_context = f"目标期刊: {journal}\n"
        if context:
            full_context += f"参考背景/意见: {context}\n"
        
        self.opt_dual_output.clear()
        self.is_processing = True
        self.status_bar.set_status("正在进行流式优化...", "warning")
        
        # 启动精确进度条
        self.precise_progress["optimize"].start(len(sections), "准备流式优化...")
        
        # 汇总所有选中的章节内容进行流式优化
        def process_sequential():
            from agents.optimizer import OptimizerAgent
            agent = OptimizerAgent(stage=stage)
            
            # 构造合并生成器，按顺序流式输出各章节
            def combined_generator():
                for i, s in enumerate(sections, 1):
                    # 更新进度条和报告区
                    self._safe_update(lambda idx=i, sec=s: self.precise_progress["optimize"].update(idx, f"正在优化: {sec}"))
                    self._safe_update(lambda sec=s: self.opt_dual_output.report_output.append_chunk(f"▶️ 正在优化章节: {sec}\n"))
                    
                    yield f"\n## {s.upper()}\n\n"
                    yield from agent.optimize_single_section_stream(s, text, full_context)
                    yield "\n\n"
            
            def on_complete(final_text):
                report = f"### 优化完成\n- 阶段: {stage}\n- 期刊: {journal}\n- 章节: {', '.join(sections)}"
                self.opt_dual_output.report_output.set_content(report)
                
                # 保存历史记录
                self.history.save_record(
                    action_type="optimize",
                    input_content=text,
                    output_content=final_text,
                    report=report,
                    metadata={'stage': stage, 'sections': sections, 'context': context}
                )
                
                self.notification.show("优化完成", "success")
                self.status_bar.set_status("优化完成", "success")
                self.precise_progress["optimize"].stop()
                self.is_processing = False
            
            def on_error(err):
                self.notification.show(f"优化失败: {err}", "error")
                self.status_bar.set_status("优化出错", "error")
                self.precise_progress["optimize"].stop(success=False)
                self.is_processing = False

            self.opt_dual_output.content_output.stream_from_generator(
                combined_generator(),
                on_complete=on_complete,
                on_error=on_error
            )

        process_sequential()
    
    def _run_dedup(self):
        """运行降重 - 支持批量处理 (P3)"""
        if not self._check_api_before_action("智能降重"):
            return
        
        files_to_process = self.dedup_file_paths if hasattr(self, 'dedup_file_paths') and self.dedup_file_paths else []
        is_batch = len(files_to_process) > 1
        
        if not files_to_process:
            text = self.dedup_input_comp.get_content()
            if not text:
                self.notification.show("请上传文件或粘贴文本", "warning")
                return
            process_queue = [(None, text)]
        else:
            process_queue = [(fp, None) for fp in files_to_process]
            
        strength = self.dedup_strength.get()
        terms_str = self.dedup_terms.get().strip()
        if "逗号分隔" in terms_str: terms_str = ""
        terms = [t.strip() for t in terms_str.split(",") if t.strip()] if terms_str else None
        
        self.dedup_dual_output.clear()
        
        def do_batch_dedup(check_cancel):
            from engines.dedup import DedupEngine
            engine = DedupEngine()
            total = len(process_queue)
            self._safe_update(lambda: self.precise_progress["dedup"].start(total, "开始批量降重..."))
            
            batch_results = []
            for i, (f_path, raw_text) in enumerate(process_queue, 1):
                if check_cancel(): return None
                fname = os.path.basename(f_path) if f_path else "粘贴的文本"
                # 修复 Lambda 闭包陷阱：使用默认参数捕获当前值
                self._safe_update(lambda idx=i, name=fname, t=total: self.precise_progress["dedup"].update(idx, f"处理中 ({idx}/{t}): {name}"))
                
                try:
                    content = raw_text
                    if f_path:
                        # 简单处理：降重引擎通常处理文本，如果是文件则尝试读取（此处简化，实际应调用 parser）
                        with open(f_path, "r", encoding="utf-8", errors="ignore") as f: content = f.read()
                    
                    if content is None: continue
                    result = engine.process(str(content), strength=int(strength), preserve_terms=terms)
                    report = engine.get_dedup_report(result)
                    
                    res_obj = {
                        'filename': fname,
                        'content': result.processed,
                        'report': f"### 文件: {fname}\n强度: {strength}\n\n{report}"
                    }
                    batch_results.append(res_obj)
                    
                    # 保存历史
                    self.history.save_record(
                        action_type="dedup",
                        input_content=str(content),
                        output_content=result.processed,
                        report=res_obj['report'],
                        metadata={'strength': strength, 'terms': terms}
                    )
                except Exception as e:
                    batch_results.append({'filename': fname, 'content': '', 'report': f"### 文件: {fname}\n❌ 失败: {e}"})
            return batch_results

        def on_complete(results):
            if results:
                if is_batch:
                    all_content = [f"--- {r['filename']} ---\n{r['content']}" for r in results]
                    all_reports = [r['report'] for r in results]
                    self.dedup_dual_output.set_content("\n\n".join(all_content), "# 批量降重报告\n\n" + "\n\n".join(all_reports))
                else:
                    self.dedup_dual_output.set_result(results[0])
                self.notification.show(f"降重完成 (共 {len(results)} 项)", "success")
            self.precise_progress["dedup"].stop()
            self.is_processing = False

        def on_error(err):
            self.notification.show(f"降重出错: {err}", "error")
            self.precise_progress["dedup"].stop(success=False)
            self.is_processing = False
            
        self.is_processing = True
        self.task_manager.submit(do_batch_dedup, on_complete=on_complete, on_error=on_error, task_name="dedup")
    
    def _run_deai(self):
        """运行降AI"""
        if not self._check_api_before_action("降AI痕迹"):
            return
        
        text = self.dedup_input.get("1.0", tk.END).strip()
        if not text:
            self.notification.show("请输入文本", "warning")
            return
        
        self._set_result(self.dedup_output, "")
        
        def do_deai(check_cancel):
            from engines.deai import DeAIEngine
            engine = DeAIEngine()
            result = engine.process(text)
            if check_cancel(): return None
            report = engine.get_report(result)
            
            # 返回结构化结果
            return {
                'content': result.processed,  # 纯净的降AI结果
                'report': f"""🤖 降AI分析报告

{'='*50}

{report}
"""
            }
        
        def on_complete(res):
            if res:
                self.dedup_dual_output.set_result(res)
                self.notification.show("消除AI痕迹完成", "success")
                self.status_bar.set_status("降AI完成", "success")
            self.precise_progress["dedup"].stop()
            self.is_processing = False

        def on_error(err):
            self._set_result(self.dedup_output, f"降AI失败: {err}")
            self.precise_progress["dedup"].stop(success=False)
            self.is_processing = False
            
        self.is_processing = True
        self.status_bar.set_status("正在消除AI痕迹...", "warning")
        task_id = self.task_manager.submit(do_deai, on_complete=on_complete, on_error=on_error, task_name="deai")
        self.precise_progress["dedup"].start(1, "正在消除AI痕迹...", on_cancel=lambda: self.task_manager.cancel(task_id))
    
    def _run_both_dedup(self):
        """运行降重+降AI"""
        if not self._check_api_before_action("深度处理"):
            return
        
        text = self.dedup_input.get("1.0", tk.END).strip()
        if not text:
            self.notification.show("请输入文本", "warning")
            return
        
        strength = self.dedup_strength.get()
        terms_str = self.dedup_terms.get().strip()
        if "逗号分隔" in terms_str:
            terms_str = ""
        terms = [t.strip() for t in terms_str.split(",") if t.strip()] if terms_str else None
        
        self._set_result(self.dedup_output, "")
        
        def do_both(check_cancel):
            from engines.dedup import DedupEngine
            from engines.deai import DeAIEngine
            
            self._safe_update(lambda: self.precise_progress["dedup"].update(1, "第1步: 智能降重..."))
            dedup_engine = DedupEngine()
            dedup_result = dedup_engine.process(text, strength=int(strength), preserve_terms=terms)
            
            if check_cancel(): return None
            
            self._safe_update(lambda: self.precise_progress["dedup"].update(2, "第2步: 消除AI痕迹..."))
            deai_engine = DeAIEngine()
            deai_result = deai_engine.process(dedup_result.processed)
            
            if check_cancel(): return None
            
            # 返回结构化结果
            return {
                'content': deai_result.processed,  # 最终处理结果
                'report': f"""⚡ 深度处理报告

处理流程: 降重 → 降AI
处理强度: {strength}/5

{'='*50}

📉 降重报告
{dedup_engine.get_dedup_report(dedup_result)}

{'='*50}

🤖 降AI报告
{deai_engine.get_report(deai_result)}
"""
            }
        
        def on_complete(res):
            if res:
                self.dedup_dual_output.set_result(res)
                
                # 保存历史记录
                self.history.save_record(
                    action_type="deep_process",
                    input_content=text,
                    output_content=res['content'],
                    report=res['report'],
                    metadata={'strength': strength, 'terms': terms}
                )
                
                self.notification.show("深度处理完成", "success")
                self.status_bar.set_status("深度处理完成", "success")
            self.precise_progress["dedup"].stop()
            self.is_processing = False

        def on_error(err):
            self._set_result(self.dedup_output, f"处理失败: {err}")
            self.precise_progress["dedup"].stop(success=False)
            self.is_processing = False
            
        self.is_processing = True
        self.status_bar.set_status("正在进行深度处理...", "warning")
        task_id = self.task_manager.submit(do_both, on_complete=on_complete, on_error=on_error, task_name="both_dedup")
        self.precise_progress["dedup"].start(2, "正在深度处理...", on_cancel=lambda: self.task_manager.cancel(task_id))
    
    def _ai_expand_keywords(self):
        """AI智能扩展关键词"""
        if not self._check_api_before_action("AI扩展关键词"):
            return
        
        query = self.search_query.get().strip()
        if not query:
            self.notification.show("请先输入初始关键词", "warning")
            return
        
        self.progress_indicators["search"].start("AI正在扩展关键词...")
        
        def do_expand(check_cancel):
            if OpenAI is None: raise ImportError("未安装 openai 库")
            from config.settings import settings
            client = OpenAI(base_url=settings.llm_api_base, api_key=settings.llm_api_key)
            
            prompt = f"作为学术研究助手，请帮我扩展以下研究主题的关键词，用于文献检索。\n\n研究主题：{query}\n\n请提供：\n1. 中文关键词扩展\n2. 英文关键词扩展\n3. 推荐的搜索组合"
            
            response = client.chat.completions.create(
                model=settings.llm_model,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.7
            )
            return response.choices[0].message.content
        
        def on_complete(result):
            self.notification.show("关键词扩展完成", "success")
            # 将结果直接显示在搜索结果区，方便用户查看和复制
            current_text = self.search_result.get("1.0", tk.END).strip()
            header = f"{'='*30} 🤖 AI 关键词扩展建议 {'='*30}\n\n研究主题：{query}\n\n{result}\n\n"
            if current_text:
                new_text = header + f"{'='*60}\n\n" + current_text
            else:
                new_text = header
            
            self._set_result(self.search_result, new_text)
            self.progress_indicators["search"].stop()
            
        def on_error(err):
            self.notification.show(f"AI扩展失败: {err}", "error")
            self.progress_indicators["search"].stop()
            
        task_id = self.task_manager.submit(do_expand, on_complete=on_complete, on_error=on_error, task_name="expand_keywords")
        self.progress_indicators["search"].start("AI正在扩展关键词...", on_cancel=lambda: self.task_manager.cancel(task_id))
    
    def _run_search(self):
        """运行学术搜索 - v2.0 使用可靠的学术API"""
        query = self.search_query.get().strip()
        if not query:
            self.notification.show("请输入搜索关键词", "warning")
            return
        
        source = self.search_source.get()
        limit = int(self.search_limit.get()) if hasattr(self, 'search_limit') else 15
        enable_ai = self.enable_ai_filter.get() if hasattr(self, 'enable_ai_filter') else False
        
        # 获取年份筛选
        year_from = None
        try:
            year_str = self.search_year_from.get().strip()
            if year_str:
                year_from = int(year_str)
        except (ValueError, AttributeError):
            pass
        
        self._set_result(self.search_result, "")
        self._safe_update(lambda: self.search_status_label.config(text="搜索中..."))
        
        def do_search(check_cancel):
            all_results = []
            errors = []
            
            # 搜索数据源
            if source in ["英文文献", "Semantic Scholar"]:
                self._safe_update(lambda: self.progress_indicators["search"].update_text("正在搜索 Semantic Scholar..."))
                try:
                    from knowledge.search.semantic_scholar import search_semantic_scholar
                    ss_results = search_semantic_scholar(query, limit=limit, year_from=year_from)
                    all_results.extend([{'title': r.title, 'authors': r.authors, 'year': r.year, 'abstract': r.abstract,
                                       'url': r.link, 'citations': r.citations, 'journal': r.venue, 'doi': r.doi,
                                       'source': 'Semantic Scholar'} for r in ss_results])
                except Exception as e: errors.append(f"SS: {e}")

            if check_cancel(): return None

            if source in ["英文文献", "OpenAlex"]:
                self._safe_update(lambda: self.progress_indicators["search"].update_text("正在搜索 OpenAlex..."))
                try:
                    from knowledge.search.openalex import search_openalex
                    oa_results = search_openalex(query, limit=limit, year_from=year_from)
                    all_results.extend([{'title': r.title, 'authors': r.authors, 'year': r.year, 'abstract': r.abstract,
                                       'url': r.link, 'citations': r.citations, 'journal': r.venue, 'doi': r.doi,
                                       'source': 'OpenAlex'} for r in oa_results])
                except Exception as e: errors.append(f"OA: {e}")

            if check_cancel(): return None

            if source in ["中文文献", "百度学术"]:
                self._safe_update(lambda: self.progress_indicators["search"].update_text("正在搜索中文文献..."))
                try:
                    from knowledge.search.cnki import search_cnki
                    cnki_results = search_cnki(query, limit=limit)
                    all_results.extend([{'title': r.title, 'authors': r.authors, 'year': r.year, 'abstract': r.abstract,
                                       'url': r.link, 'citations': r.citations, 'journal': r.source, 'doi': '',
                                       'source': r.database} for r in cnki_results])
                except Exception as e: errors.append(f"CNKI: {e}")

            if check_cancel(): return None

            # 筛选逻辑
            if all_results:
                try:
                    from knowledge.search.journal_rank import enrich_with_rank_info, filter_by_quality
                    all_results = enrich_with_rank_info(all_results)
                    if self.filter_cssci.get() or self.filter_ssci.get():
                        all_results = filter_by_quality(all_results, require_cssci=self.filter_cssci.get(), require_ssci=self.filter_ssci.get())
                except: pass

            if not all_results:
                return {"error": "未找到相关文献。\n\n" + "\n".join(errors)}

            # 去重和排序
            seen = set()
            unique = []
            for p in all_results:
                if p['title'].lower() not in seen:
                    seen.add(p['title'].lower())
                    unique.append(p)
            
            if check_cancel(): return None
            
            if enable_ai and len(unique) > limit:
                self._safe_update(lambda: self.progress_indicators["search"].update_text("AI智能筛选中..."))
                unique = self._ai_filter_papers(query, unique, limit)
                
            unique.sort(key=lambda x: x.get('citations', 0) or 0, reverse=True)
            return unique

        def on_complete(results):
            if isinstance(results, dict) and "error" in results:
                self._set_result(self.search_result, str(results["error"]))
                self.search_status_label.config(text="未找到结果")
            elif isinstance(results, list):
                formatted = self._format_search_results(results, enable_ai)
                # 使用 DualOutputFrame 显示结果
                self.search_dual_output.set_content(formatted, f"🔍 搜索报告\n\n关键词: {query}\n数据源: {source}\n结果数量: {len(results)}")
                
                # 保存历史记录
                self.history.save_record(
                    action_type="search",
                    input_content=query,
                    output_content=formatted,
                    report=f"🔍 搜索报告\n数据源: {source}",
                    metadata={'query': query, 'source': source, 'results_count': len(results)}
                )
                
                self.search_status_label.config(text=f"共 {len(results)} 篇文献")
                self.last_search_results = results
                self.notification.show("搜索完成", "success")
                self.status_bar.set_status(f"搜索完成，找到 {len(results)} 篇文献", "success")
            self.precise_progress["search"].stop()
            self.is_processing = False

        def on_error(err):
            self._set_result(self.search_result, f"搜索失败: {err}")
            self.progress_indicators["search"].stop()
            self.is_processing = False
            
        self.is_processing = True
        self.status_bar.set_status(f"正在搜索 {source}...", "warning")
        task_id = self.task_manager.submit(do_search, on_complete=on_complete, on_error=on_error, task_name="search")
        self.precise_progress["search"].start(1, f"正在搜索 {source}...", on_cancel=lambda: self.task_manager.cancel(task_id))
    
    def _generate_literature_review(self):
        """基于搜索结果生成文献综述"""
        if not self._check_api_before_action("生成文献综述"):
            return
        
        if not hasattr(self, 'last_search_results') or not self.last_search_results:
            self.notification.show("请先搜索文献", "warning")
            return
        
        def generate_stream():
            if OpenAI is None: raise ImportError("未安装 openai 库")
            from config.settings import settings
            client = OpenAI(base_url=settings.llm_api_base, api_key=settings.llm_api_key)
            
            papers_text = ""
            for i, p in enumerate(self.last_search_results[:15], 1):
                papers_text += f"{i}. {p.get('title')} ({p.get('authors')}, {p.get('year')})\n摘要：{p.get('abstract', '')[:300]}\n\n"
            
            prompt = f"请基于以下学术文献，生成一段学术论文风格的文献综述（约500-800字）。\n\n要求：1. 客观严谨 2. 归纳对比 3. 正确引用 4. 指出共识分歧\n\n文献列表：\n{papers_text}"
            
            response = client.chat.completions.create(
                model=settings.llm_model,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.7,
                max_tokens=2000,
                stream=True
            )
            for chunk in response:
                if chunk.choices and chunk.choices[0].delta.content:
                    yield chunk.choices[0].delta.content
        
        # 清空并准备流式输出到报告区
        current_content = self.search_dual_output.get_content()
        self.search_dual_output.report_output.start_streaming("正在生成文献综述...")
        self.notebook = self.search_dual_output.notebook
        self.notebook.select(1) # 切换到报告页
        
        def on_complete(review):
            if review:
                current_report = self.search_dual_output.get_report()
                # 格式化最终报告
                final_report = f"{'='*60}\n📝 AI 文献综述\n{'='*60}\n\n{review}\n\n{current_report}"
                self.search_dual_output.report_output.set_content(final_report)
                
                self.notification.show("文献综述生成完成", "success")
                self.status_bar.set_status("文献综述生成完成", "success")
            self.precise_progress["search"].stop()
            self.is_processing = False

        def on_error(err):
            self.notification.show(f"生成失败: {err}", "error")
            self.progress_indicators["search"].stop()
            self.is_processing = False
            
        self.is_processing = True
        self.status_bar.set_status("正在生成文献综述...", "warning")
        
        # 使用流式组件执行生成器
        self.search_dual_output.report_output.stream_from_generator(
            generate_stream(),
            on_complete=on_complete,
            on_error=on_error
        )
        self.precise_progress["search"].start(1, "正在生成文献综述...")
    
    def _generate_citations(self):
        """生成引用格式"""
        if not hasattr(self, 'last_search_results') or not self.last_search_results:
            self.notification.show("请先搜索文献", "warning")
            return
        
        # 创建选择窗口
        cite_window = tk.Toplevel(self.root)
        cite_window.title("选择引用格式")
        cite_window.geometry("600x500")
        cite_window.configure(bg=ModernStyle.BG_MAIN)
        cite_window.transient(self.root)
        
        # 居中显示
        cite_window.update_idletasks()
        x = (cite_window.winfo_screenwidth() - 600) // 2
        y = (cite_window.winfo_screenheight() - 500) // 2
        cite_window.geometry(f"+{x}+{y}")
        
        content = tk.Frame(cite_window, bg=ModernStyle.BG_MAIN, padx=25, pady=20)
        content.pack(fill=tk.BOTH, expand=True)
        
        tk.Label(
            content,
            text="选择引用格式",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_LG, "bold"),
            bg=ModernStyle.BG_MAIN
        ).pack(anchor="w", pady=(0, 15))
        
        style_var = tk.StringVar(value="apa")
        styles = [
            ("APA 格式", "apa"),
            ("GB/T 7714 格式（中国国标）", "gb"),
            ("MLA 格式", "mla"),
            ("Chicago 格式", "chicago")
        ]
        
        for text, value in styles:
            tk.Radiobutton(
                content,
                text=text,
                variable=style_var,
                value=value,
                bg=ModernStyle.BG_MAIN,
                font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_SM)
            ).pack(anchor="w", pady=3)
        
        # 引用预览区
        tk.Label(
            content,
            text="引用预览：",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_MD, "bold"),
            bg=ModernStyle.BG_MAIN
        ).pack(anchor="w", pady=(20, 10))
        
        preview_text = scrolledtext.ScrolledText(
            content,
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_SM),
            height=15,
            wrap=tk.WORD,
            bg=ModernStyle.BG_SECONDARY
        )
        preview_text.pack(fill=tk.BOTH, expand=True)
        
        def update_preview(*args):
            style = style_var.get()
            citations = []
            
            for i, p in enumerate(self.last_search_results[:20], 1):
                authors = p.get('authors', '未知作者')
                year = p.get('year', '')
                title = p.get('title', '无标题')
                journal = p.get('journal', '')
                doi = p.get('doi', '')
                
                if style == "apa":
                    cite = f"{authors} ({year}). {title}."
                    if journal:
                        cite += f" {journal}."
                    if doi:
                        cite += f" https://doi.org/{doi}"
                elif style == "gb":
                    cite = f"[{i}] {authors}. {title}[J]. {journal}, {year}."
                elif style == "mla":
                    cite = f'{authors}. "{title}." {journal}, {year}.'
                elif style == "chicago":
                    cite = f'{authors}. "{title}." {journal} ({year}).'
                    if doi:
                        cite += f" https://doi.org/{doi}."
                else:
                    cite = f"{authors} ({year}). {title}. {journal}."
                
                citations.append(cite)
            
            preview_text.delete("1.0", tk.END)
            preview_text.insert("1.0", "\n\n".join(citations))
        
        style_var.trace("w", update_preview)
        update_preview()
        
        # 按钮
        btn_frame = tk.Frame(content, bg=ModernStyle.BG_MAIN)
        btn_frame.pack(fill=tk.X, pady=(15, 0))
        
        def copy_citations():
            self.root.clipboard_clear()
            self.root.clipboard_append(preview_text.get("1.0", tk.END).strip())
            self.notification.show("引用已复制到剪贴板", "success")
        
        ModernButton(
            btn_frame,
            text="📋 复制引用",
            command=copy_citations,
            width=120,
            height=40
        ).pack(side=tk.LEFT)
        
        ModernButton(
            btn_frame,
            text="关闭",
            command=cite_window.destroy,
            width=100,
            height=40,
            bg_color=ModernStyle.BG_SECONDARY,
            hover_color=ModernStyle.BG_HOVER,
            text_color=ModernStyle.TEXT_PRIMARY
        ).pack(side=tk.LEFT, padx=15)
    
    def _filter_by_journal_rank(self, papers: list, source_type: str, show_rank: bool) -> list:
        """根据期刊级别过滤论文
        
        Args:
            papers: 论文列表
            source_type: "chinese" 或 "english"
            show_rank: 是否显示期刊级别信息
        
        Returns:
            过滤后的论文列表
        """
        try:
            from knowledge.search.journal_rank import check_journal_rank, is_high_quality_journal, format_rank_info
            
            filtered = []
            for paper in papers:
                journal = paper.get("journal", "")
                if not journal:
                    # 如果没有期刊信息，暂时保留
                    filtered.append(paper)
                    continue
                
                rank = check_journal_rank(journal)
                
                if is_high_quality_journal(rank, source_type):
                    if show_rank and rank:
                        paper["rank_info"] = format_rank_info(rank)
                    filtered.append(paper)
            
            # 如果所有论文都被过滤掉，返回前5条原始结果
            return filtered if filtered else papers[:5]
            
        except Exception as e:
            print(f"期刊级别查询失败: {e}")
            return papers
    
    def _ai_filter_papers(self, query: str, papers: list, top_k: int) -> list:
        """AI智能筛选文献"""
        try:
            from openai import OpenAI
            from config.settings import settings
            
            client = OpenAI(base_url=settings.llm_api_base, api_key=settings.llm_api_key)
            
            # 构建文献摘要
            papers_text = ""
            for i, p in enumerate(papers[:30], 1):  # 最多30篇供筛选
                title = p.get('title', '无标题')
                abstract = p.get('abstract', p.get('snippet', '无摘要'))[:150]
                papers_text += f"{i}. {title}\n   摘要：{abstract}\n\n"
            
            prompt = f"""作为学术研究助手，请从以下文献中筛选出与研究主题最相关的 {top_k} 篇。

研究主题：{query}

文献列表：
{papers_text}

请仅返回最相关的文献序号（用逗号分隔，如：1,5,8），从最相关到较相关排序。"""

            response = client.chat.completions.create(
                model=settings.llm_model,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.3,
                max_tokens=100
            )
            
            content = response.choices[0].message.content
            if content is None:
                return papers[:top_k]
            selected = content.strip()
            indices = [int(x.strip()) - 1 for x in selected.split(',') if x.strip().isdigit()]
            
            return [papers[i] for i in indices if 0 <= i < len(papers)]
            
        except Exception:
            return papers[:top_k]
    
    def _format_search_results(self, results: list, ai_filtered: bool) -> str:
        """格式化搜索结果"""
        if not results:
            return "未找到相关文献"
        
        output = []
        output.append(f"{'='*60}")
        output.append(f"📚 检索结果：共找到 {len(results)} 篇文献" + (" (AI智能筛选)" if ai_filtered else ""))
        output.append(f"{'='*60}\n")
        
        for i, paper in enumerate(results, 1):
            source = paper.get('source', '未知来源')
            title = paper.get('title', '无标题')
            authors = paper.get('authors', '未知作者')
            year = paper.get('year', '未知年份')
            journal = paper.get('journal', paper.get('venue', ''))
            citations = paper.get('citations', 0)
            abstract = paper.get('abstract', paper.get('snippet', '无摘要'))
            url = paper.get('url', '')
            
            output.append(f"【{i}】{title}")
            output.append(f"    来源: {source}")
            output.append(f"    作者: {authors}")
            output.append(f"    发表: {year}" + (f" | {journal}" if journal else ""))
            
            # 显示期刊级别
            rank_info = paper.get("rank_info", "")
            if rank_info:
                output.append(f"    📊 级别: {rank_info}")
            
            if citations:
                output.append(f"    引用: {citations}")
            output.append(f"    摘要: {abstract[:250]}...")
            if url:
                output.append(f"    链接: {url}")
            output.append("")
        
        output.append(f"\n{'='*60}")
        output.append("💡 提示：点击「深度优化」→「引用文献」将搜索结果融入论文")
        output.append("💡 提示：点击「退修助手」→「找支撑文献」获取审稿回应所需参考")
        
        return "\n".join(output)
    
    def _recommend_literature(self):
        """根据论文内容智能推荐文献"""
        content = self.diag_text.get("1.0", tk.END).strip()
        if not content:
            self.notification.show("请先输入论文内容", "warning")
            return
        
        def do_recommend(check_cancel):
            if OpenAI is None: raise ImportError("未安装 openai 库")
            from config.settings import settings
            client = OpenAI(base_url=settings.llm_api_base, api_key=settings.llm_api_key)
            
            prompt = f"请分析以下论文内容，提取3-5个核心研究关键词用于文献检索：\n\n{content[:2000]}\n\n仅返回关键词，逗号分隔。"
            response = client.chat.completions.create(
                model=settings.llm_model,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.5,
                max_tokens=100
            )
            return response.choices[0].message.content
        
        def on_complete(keywords):
            if keywords:
                self.search_query.delete(0, tk.END)
                self.search_query.insert(0, keywords.strip())
                self._show_page("search")
                self.root.after(500, self._run_search)
            self.progress_indicators["diagnose"].stop()
            self.is_processing = False

        def on_error(err):
            self.notification.show(f"推荐失败: {err}", "error")
            self.progress_indicators["diagnose"].stop()
            self.is_processing = False
            
        self.is_processing = True
        task_id = self.task_manager.submit(do_recommend, on_complete=on_complete, on_error=on_error, task_name="recommend")
        self.progress_indicators["diagnose"].start("AI正在推荐文献...", on_cancel=lambda: self.task_manager.cancel(task_id))
    
    def _find_supporting_literature(self):
        """根据审稿意见找支撑文献"""
        comments = self.rev_comments.get("1.0", tk.END).strip()
        if not comments:
            self.notification.show("请先输入审稿意见", "warning")
            return
        
        def do_find(check_cancel):
            if OpenAI is None: raise ImportError("未安装 openai 库")
            from config.settings import settings
            client = OpenAI(base_url=settings.llm_api_base, api_key=settings.llm_api_key)
            
            prompt = f"请分析以下审稿意见，提取关键词用于查找支撑文献：\n\n{comments[:1500]}\n\n仅返回关键词，逗号分隔。"
            response = client.chat.completions.create(
                model=settings.llm_model,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.5,
                max_tokens=100
            )
            return response.choices[0].message.content
        
        def on_complete(keywords):
            if keywords:
                self.search_query.delete(0, tk.END)
                self.search_query.insert(0, keywords.strip())
                self._show_page("search")
                self.root.after(500, self._run_search)
            self.progress_indicators["revision"].stop()
            self.is_processing = False

        def on_error(err):
            self.notification.show(f"查找失败: {err}", "error")
            self.progress_indicators["revision"].stop()
            self.is_processing = False
            
        self.is_processing = True
        task_id = self.task_manager.submit(do_find, on_complete=on_complete, on_error=on_error, task_name="find_lit")
        self.progress_indicators["revision"].start("AI正在提取文献关键词...", on_cancel=lambda: self.task_manager.cancel(task_id))
    
    def _run_revision(self):
        """运行退修处理"""
        if not self._check_api_before_action("退修助手"):
            return
        
        comments = self.rev_comments.get("1.0", tk.END).strip()
        if not comments:
            self.notification.show("请粘贴审稿意见", "warning")
            return
        
        summary = self.rev_summary.get("1.0", tk.END).strip() or None
        
        self.rev_dual_output.clear()
        self.is_processing = True
        self.status_bar.set_status("正在生成退修回应...", "warning")
        
        from agents.revision import RevisionAgent
        agent = RevisionAgent()
        
        # 使用流式输出
        self.rev_dual_output.content_output.start_streaming("正在处理审稿意见...")
        self.precise_progress["revision"].start(2, "正在解析审稿意见...")
        
        def on_complete(final_letter):
            self._safe_update(lambda: self.precise_progress["revision"].update(1, "生成建议信完成，正在生成分析报告..."))
            # 核心内容（回复信）生成完成后，启动后台任务生成分析报告
            def get_report_task(check_cancel):
                return agent.process_comments(comments, summary)
            
            def on_report_ready(result):
                formatted = agent.format_result(result)
                report_text = f"""📝 审稿意见分析报告

{'='*50}

{formatted}

{'='*50}

💡 修改建议汇总
- 请根据上述分析逐条修改论文
- 建议使用「查找文献」功能获取支撑材料
"""
                self.rev_dual_output.report_output.set_content(report_text)
                
                # 保存历史记录
                self.history.save_record(
                    action_type="revision",
                    input_content=comments,
                    output_content=final_letter,
                    report=report_text
                )
                
                self.notification.show("退修建议生成完成", "success")
                self.status_bar.set_status("退修建议生成完成", "success")
                self.precise_progress["revision"].stop()
                self.is_processing = False
            
            # 提交任务以获取结构化报告
            self.task_manager.submit(get_report_task, on_complete=on_report_ready)

        def on_error(err):
            self.notification.show(f"处理失败: {err}", "error")
            self.rev_dual_output.content_output.end_streaming(False)
            self.precise_progress["revision"].stop(success=False)
            self.is_processing = False
            
        self.rev_dual_output.content_output.stream_from_generator(
            agent.process_comments_stream(comments, summary),
            on_complete=on_complete,
            on_error=on_error
        )
    
    def _load_settings(self):
        """加载设置"""
        try:
            from config.settings import settings
            self.setting_llm_base.delete(0, tk.END)
            self.setting_llm_base.insert(0, settings.llm_api_base or "")
            self.setting_llm_key.delete(0, tk.END)
            self.setting_llm_key.insert(0, settings.llm_api_key or "")
            self.setting_llm_model.set(settings.llm_model or "gpt-4o-mini")
            
            # 安全访问嵌入模型配置控件
            if hasattr(self, 'setting_embed_base') and hasattr(self.setting_embed_base, 'winfo_exists') and self.setting_embed_base.winfo_exists():
                self.setting_embed_base.delete(0, tk.END)
                self.setting_embed_base.insert(0, settings.embedding_api_base or "")
            if hasattr(self, 'setting_embed_key') and hasattr(self.setting_embed_key, 'winfo_exists') and self.setting_embed_key.winfo_exists():
                self.setting_embed_key.delete(0, tk.END)
                self.setting_embed_key.insert(0, settings.embedding_api_key or "")
            if hasattr(self, 'setting_embed_model'):
                self.setting_embed_model.set(settings.embedding_model or "text-embedding-3-small")
            
            if settings.llm_api_key:
                self.llm_status.config(text="● 已配置", fg=ModernStyle.SUCCESS)
        except Exception:
            pass
    

    def _save_settings(self):
        """保存设置"""
        try:
            env_path = BASE_DIR / ".env"
            
            if self.use_same_api.get():
                embed_base = self.setting_llm_base.get()
                embed_key = self.setting_llm_key.get()
            else:
                # 安全访问嵌入模型配置控件
                embed_base = self.setting_embed_base.get() if (hasattr(self, 'setting_embed_base') and hasattr(self.setting_embed_base, 'winfo_exists') and self.setting_embed_base.winfo_exists()) else self.setting_llm_base.get()
                embed_key = self.setting_embed_key.get() if (hasattr(self, 'setting_embed_key') and hasattr(self.setting_embed_key, 'winfo_exists') and self.setting_embed_key.winfo_exists()) else self.setting_llm_key.get()
            
            # 获取存储目录配置
            data_dir = self.setting_data_dir.get().strip() if hasattr(self, 'setting_data_dir') else ""
            workspace_dir = self.setting_workspace_dir.get().strip() if hasattr(self, 'setting_workspace_dir') else ""
            
            lines = [
                f"# EconPaper Pro 配置",
                f"",
                f"# 语言模型 (LLM) 配置",
                f"LLM_API_BASE={self.setting_llm_base.get()}",
                f"LLM_API_KEY={self.setting_llm_key.get()}",
                f"LLM_MODEL={self.setting_llm_model.get()}",
                f"",
                f"# 嵌入模型 (Embedding) 配置",
                f"EMBEDDING_API_BASE={embed_base}",
                f"EMBEDDING_API_KEY={embed_key}",
                f"EMBEDDING_MODEL={self.setting_embed_model.get()}",
                f"",
                f"# 存储目录配置 (避免占用C盘)",
                f"DATA_DIR={data_dir}",
                f"WORKSPACE_DIR={workspace_dir}",
            ]
            
            with open(env_path, "w", encoding="utf-8") as f:
                f.write("\n".join(lines))
            
            self.llm_status.config(text="● 已配置", fg=ModernStyle.SUCCESS)
            self.notification.show("配置已保存！部分设置重启生效。", "success")
            
        except Exception as e:
            self.notification.show(f"保存失败: {e}", "error")


def main():
    """主函数"""
    root = tk.Tk()
    
    # 设置 DPI 感知（Windows）
    try:
        from ctypes import windll
        windll.shcore.SetProcessDpiAwareness(1)
    except Exception:
        pass
    
    app = EconPaperApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
